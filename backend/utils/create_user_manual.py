"""
Скрипт для создания инструкции по работе с платформой AGB Passports
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_user_manual():
    """Создает DOCX файл с инструкцией по работе с платформой"""
    
    # Создаем документ
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Заголовок документа
    title = doc.add_heading('Инструкция по работе с платформой AGB Passports', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Подзаголовок
    subtitle = doc.add_paragraph('Система управления паспортами и наклейками для бурового инструмента')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(14)
    subtitle_format.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()  # Пустая строка
    
    # Содержание
    doc.add_heading('Содержание', 1)
    doc.add_paragraph('1. Вход в систему')
    doc.add_paragraph('2. Создание паспортов')
    doc.add_paragraph('3. Работа с архивом паспортов')
    doc.add_paragraph('4. Работа с архивом наклеек')
    doc.add_paragraph('5. Управление номенклатурой (для администраторов)')
    doc.add_paragraph('6. Управление пользователями (для администраторов)')
    doc.add_paragraph('7. Экспорт данных')
    doc.add_paragraph('8. Редактор шаблонов (для администраторов)')
    
    doc.add_page_break()
    
    # Раздел 1: Вход в систему
    doc.add_heading('1. Вход в систему', 1)
    
    doc.add_paragraph('Для начала работы с платформой необходимо выполнить вход в систему.')
    
    doc.add_paragraph('Шаги:', style='List Number')
    doc.add_paragraph('Откройте браузер и перейдите по адресу платформы', style='List Bullet 2')
    doc.add_paragraph('На странице входа введите ваши учетные данные:', style='List Bullet 2')
    doc.add_paragraph('   • Имя пользователя (username)', style='List Bullet 3')
    doc.add_paragraph('   • Пароль', style='List Bullet 3')
    doc.add_paragraph('Нажмите кнопку "Войти"', style='List Bullet 2')
    
    # Место для скриншота
    doc.add_paragraph()
    screenshot_para = doc.add_paragraph('[СКРИНШОТ: Страница входа в систему]')
    screenshot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    screenshot_para.runs[0].font.italic = True
    screenshot_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    
    doc.add_paragraph()
    doc.add_paragraph('После успешного входа вы будете перенаправлены на главную страницу платформы.')
    
    doc.add_page_break()
    
    # Раздел 2: Создание паспортов
    doc.add_heading('2. Создание паспортов', 1)
    
    doc.add_paragraph('Основная функция платформы - создание паспортов для бурового инструмента.')
    
    doc.add_heading('2.1. Выбор номенклатуры', 2)
    doc.add_paragraph('Для создания паспорта необходимо:')
    doc.add_paragraph('1. Перейти в раздел "Создание паспорта" (левое меню)')
    doc.add_paragraph('2. В поле поиска ввести название или артикул номенклатуры')
    doc.add_paragraph('3. Выбрать нужную номенклатуру из списка')
    
    doc.add_paragraph()
    screenshot_para = doc.add_paragraph('[СКРИНШОТ: Выбор номенклатуры]')
    screenshot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    screenshot_para.runs[0].font.italic = True
    screenshot_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    
    doc.add_heading('2.2. Указание количества и номера заказа', 2)
    doc.add_paragraph('После выбора номенклатуры:')
    doc.add_paragraph('1. Укажите количество паспортов для создания')
    doc.add_paragraph('2. (Опционально) Укажите номер заказа')
    doc.add_paragraph('3. Нажмите кнопку "Создать паспорта"')
    
    doc.add_paragraph()
    screenshot_para = doc.add_paragraph('[СКРИНШОТ: Форма создания паспортов]')
    screenshot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    screenshot_para.runs[0].font.italic = True
    screenshot_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    
    doc.add_heading('2.3. Результат создания', 2)
    doc.add_paragraph('После успешного создания паспортов:')
    doc.add_paragraph('• Система автоматически сгенерирует номера паспортов')
    doc.add_paragraph('• Отобразится список созданных паспортов')
    doc.add_paragraph('• Будет доступна возможность экспорта паспортов в PDF или Excel')
    
    doc.add_paragraph()
    screenshot_para = doc.add_paragraph('[СКРИНШОТ: Список созданных паспортов]')
    screenshot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    screenshot_para.runs[0].font.italic = True
    screenshot_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    
    doc.add_page_break()
    
    # Раздел 3: Работа с архивом паспортов
    doc.add_heading('3. Работа с архивом паспортов', 1)
    
    doc.add_paragraph('В разделе "Архив паспортов" можно просматривать, искать и экспортировать созданные паспорта.')
    
    doc.add_heading('3.1. Просмотр паспортов', 2)
    doc.add_paragraph('• Паспорта отображаются в виде таблицы с пагинацией')
    doc.add_paragraph('• Для обычных пользователей отображаются только их собственные паспорта')
    doc.add_paragraph('• Администраторы видят все паспорта в системе')
    
    doc.add_paragraph()
    screenshot_para = doc.add_paragraph('[СКРИНШОТ: Архив паспортов]')
    screenshot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    screenshot_para.runs[0].font.italic = True
    screenshot_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    
    doc.add_heading('3.2. Поиск и фильтрация', 2)
    doc.add_paragraph('Для поиска паспортов:')
    doc.add_paragraph('1. Используйте поле поиска для ввода номера паспорта или артикула')
    doc.add_paragraph('2. Применяйте фильтры по типу продукта, матрице или статусу')
    doc.add_paragraph('3. Результаты обновляются автоматически')
    
    doc.add_paragraph()
    screenshot_para = doc.add_paragraph('[СКРИНШОТ: Поиск и фильтры]')
    screenshot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    screenshot_para.runs[0].font.italic = True
    screenshot_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    
    doc.add_heading('3.3. Экспорт паспортов', 2)
    doc.add_paragraph('Для экспорта паспортов:')
    doc.add_paragraph('1. Выберите нужные паспорта, установив флажки в таблице')
    doc.add_paragraph('2. Нажмите кнопку "Экспорт выбранных"')
    doc.add_paragraph('3. Выберите формат экспорта:')
    doc.add_paragraph('   • PDF - для печати паспортов')
    doc.add_paragraph('   • Excel - для работы с данными в таблицах')
    
    doc.add_paragraph()
    screenshot_para = doc.add_paragraph('[СКРИНШОТ: Экспорт паспортов]')
    screenshot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    screenshot_para.runs[0].font.italic = True
    screenshot_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    
    doc.add_page_break()
    
    # Раздел 4: Работа с архивом наклеек
    doc.add_heading('4. Работа с архивом наклеек', 1)
    
    doc.add_paragraph('В разделе "Архив наклеек" можно создавать и экспортировать наклейки для паспортов.')
    
    doc.add_heading('4.1. Выбор паспортов для наклеек', 2)
    doc.add_paragraph('1. Перейдите в раздел "Архив наклеек"')
    doc.add_paragraph('2. Используйте поиск для нахождения нужных паспортов')
    doc.add_paragraph('3. Выберите паспорта, для которых нужно создать наклейки')
    
    doc.add_paragraph()
    screenshot_para = doc.add_paragraph('[СКРИНШОТ: Выбор паспортов для наклеек]')
    screenshot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    screenshot_para.runs[0].font.italic = True
    screenshot_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    
    doc.add_heading('4.2. Экспорт наклеек', 2)
    doc.add_paragraph('После выбора паспортов:')
    doc.add_paragraph('1. Нажмите кнопку "Экспорт наклеек"')
    doc.add_paragraph('2. Выберите формат:')
    doc.add_paragraph('   • DOCX - наклейки в формате Word (6 наклеек на странице, 2 столбца × 3 строки)')
    doc.add_paragraph('   • PDF - наклейки в формате PDF (8 наклеек на странице)')
    doc.add_paragraph('3. Файл будет автоматически загружен')
    
    doc.add_paragraph()
    doc.add_paragraph('Примечание: Наклейки содержат:')
    doc.add_paragraph('• Логотип компании')
    doc.add_paragraph('• Название номенклатуры')
    doc.add_paragraph('• Серийный номер паспорта')
    doc.add_paragraph('• Артикул и матрицу')
    doc.add_paragraph('• Штрихкоды для артикула и серийного номера')
    
    doc.add_paragraph()
    screenshot_para = doc.add_paragraph('[СКРИНШОТ: Экспорт наклеек]')
    screenshot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    screenshot_para.runs[0].font.italic = True
    screenshot_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    
    doc.add_page_break()
    
    # Раздел 5: Управление номенклатурой (для админов)
    doc.add_heading('5. Управление номенклатурой (для администраторов)', 1)
    
    doc.add_paragraph('Администраторы могут управлять справочником номенклатуры.')
    
    doc.add_heading('5.1. Просмотр номенклатуры', 2)
    doc.add_paragraph('1. Перейдите в раздел "Номенклатура" (доступен только администраторам)')
    doc.add_paragraph('2. Просмотрите список всей номенклатуры в системе')
    doc.add_paragraph('3. Используйте поиск для быстрого нахождения нужных позиций')
    
    doc.add_paragraph()
    screenshot_para = doc.add_paragraph('[СКРИНШОТ: Список номенклатуры]')
    screenshot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    screenshot_para.runs[0].font.italic = True
    screenshot_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    
    doc.add_heading('5.2. Добавление номенклатуры', 2)
    doc.add_paragraph('Для добавления новой номенклатуры:')
    doc.add_paragraph('1. Нажмите кнопку "Добавить номенклатуру"')
    doc.add_paragraph('2. Заполните форму:')
    doc.add_paragraph('   • Код 1С')
    doc.add_paragraph('   • Название')
    doc.add_paragraph('   • Артикул')
    doc.add_paragraph('   • Матрица')
    doc.add_paragraph('   • Высота/глубина бурения')
    doc.add_paragraph('   • Промывочные отверстия')
    doc.add_paragraph('   • Тип продукта')
    doc.add_paragraph('3. Нажмите "Сохранить"')
    
    doc.add_paragraph()
    screenshot_para = doc.add_paragraph('[СКРИНШОТ: Форма добавления номенклатуры]')
    screenshot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    screenshot_para.runs[0].font.italic = True
    screenshot_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    
    doc.add_heading('5.3. Редактирование номенклатуры', 2)
    doc.add_paragraph('1. Найдите нужную номенклатуру в списке')
    doc.add_paragraph('2. Нажмите кнопку "Редактировать"')
    doc.add_paragraph('3. Внесите необходимые изменения')
    doc.add_paragraph('4. Нажмите "Сохранить"')
    
    doc.add_heading('5.4. Деактивация номенклатуры', 2)
    doc.add_paragraph('Для деактивации номенклатуры (она перестанет отображаться при создании паспортов):')
    doc.add_paragraph('1. Откройте форму редактирования')
    doc.add_paragraph('2. Снимите флажок "Активна"')
    doc.add_paragraph('3. Сохраните изменения')
    
    doc.add_page_break()
    
    # Раздел 6: Управление пользователями (для админов)
    doc.add_heading('6. Управление пользователями (для администраторов)', 1)
    
    doc.add_paragraph('Администраторы могут управлять пользователями системы.')
    
    doc.add_heading('6.1. Просмотр пользователей', 2)
    doc.add_paragraph('1. Перейдите в раздел "Пользователи" (доступен только администраторам)')
    doc.add_paragraph('2. Просмотрите список всех пользователей системы')
    doc.add_paragraph('3. Для каждого пользователя отображается:')
    doc.add_paragraph('   • Имя пользователя')
    doc.add_paragraph('   • Email')
    doc.add_paragraph('   • Полное имя')
    doc.add_paragraph('   • Роль (admin/user)')
    doc.add_paragraph('   • Статус активности')
    doc.add_paragraph('   • Дата последнего входа')
    
    doc.add_paragraph()
    screenshot_para = doc.add_paragraph('[СКРИНШОТ: Список пользователей]')
    screenshot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    screenshot_para.runs[0].font.italic = True
    screenshot_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    
    doc.add_heading('6.2. Создание пользователя', 2)
    doc.add_paragraph('Для создания нового пользователя:')
    doc.add_paragraph('1. Нажмите кнопку "Создать пользователя"')
    doc.add_paragraph('2. Заполните форму:')
    doc.add_paragraph('   • Имя пользователя (обязательно)')
    doc.add_paragraph('   • Email (обязательно)')
    doc.add_paragraph('   • Полное имя')
    doc.add_paragraph('   • Пароль (обязательно)')
    doc.add_paragraph('   • Роль (admin или user)')
    doc.add_paragraph('3. Нажмите "Создать"')
    
    doc.add_paragraph()
    screenshot_para = doc.add_paragraph('[СКРИНШОТ: Форма создания пользователя]')
    screenshot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    screenshot_para.runs[0].font.italic = True
    screenshot_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    
    doc.add_heading('6.3. Редактирование пользователя', 2)
    doc.add_paragraph('1. Найдите нужного пользователя в списке')
    doc.add_paragraph('2. Нажмите кнопку "Редактировать"')
    doc.add_paragraph('3. Внесите необходимые изменения')
    doc.add_paragraph('4. При необходимости измените пароль')
    doc.add_paragraph('5. Нажмите "Сохранить"')
    
    doc.add_heading('6.4. Блокировка пользователя', 2)
    doc.add_paragraph('Для блокировки доступа пользователя:')
    doc.add_paragraph('1. Откройте форму редактирования пользователя')
    doc.add_paragraph('2. Снимите флажок "Активен"')
    doc.add_paragraph('3. Сохраните изменения')
    doc.add_paragraph('После этого пользователь не сможет войти в систему.')
    
    doc.add_page_break()
    
    # Раздел 7: Экспорт данных
    doc.add_heading('7. Экспорт данных', 1)
    
    doc.add_paragraph('Платформа поддерживает экспорт данных в различных форматах.')
    
    doc.add_heading('7.1. Экспорт паспортов в PDF', 2)
    doc.add_paragraph('Формат PDF используется для печати паспортов.')
    doc.add_paragraph('Характеристики:')
    doc.add_paragraph('• Каждый паспорт на отдельной странице')
    doc.add_paragraph('• Полная информация о паспорте и номенклатуре')
    doc.add_paragraph('• Готов к печати')
    
    doc.add_heading('7.2. Экспорт паспортов в Excel', 2)
    doc.add_paragraph('Формат Excel используется для работы с данными в таблицах.')
    doc.add_paragraph('Характеристики:')
    doc.add_paragraph('• Все паспорта в одной таблице')
    doc.add_paragraph('• Столбцы: номер паспорта, номенклатура, артикул, дата создания и др.')
    doc.add_paragraph('• Удобен для анализа и фильтрации данных')
    
    doc.add_heading('7.3. Экспорт наклеек в DOCX', 2)
    doc.add_paragraph('Формат DOCX используется для печати наклеек.')
    doc.add_paragraph('Характеристики:')
    doc.add_paragraph('• 6 наклеек на странице (2 столбца × 3 строки)')
    doc.add_paragraph('• Каждая наклейка содержит полную информацию')
    doc.add_paragraph('• Включает логотип и штрихкоды')
    doc.add_paragraph('• Готов к печати на самоклеящихся этикетках')
    
    doc.add_heading('7.4. Экспорт наклеек в PDF', 2)
    doc.add_paragraph('Альтернативный формат для печати наклеек.')
    doc.add_paragraph('Характеристики:')
    doc.add_paragraph('• 8 наклеек на странице')
    doc.add_paragraph('• Компактное размещение')
    doc.add_paragraph('• Включает всю необходимую информацию')
    
    doc.add_page_break()
    
    # Раздел 8: Редактор шаблонов
    doc.add_heading('8. Редактор шаблонов (для администраторов)', 1)
    
    doc.add_paragraph('Администраторы могут редактировать шаблоны для паспортов и наклеек.')
    
    doc.add_heading('8.1. Редактор шаблона паспорта', 2)
    doc.add_paragraph('1. Перейдите в раздел "Редактор шаблонов"')
    doc.add_paragraph('2. Выберите вкладку "Шаблон паспорта"')
    doc.add_paragraph('3. Загрузите существующий шаблон или создайте новый')
    doc.add_paragraph('4. Используйте плейсхолдеры для вставки данных:')
    doc.add_paragraph('   • {{nomenclature_name}} - название номенклатуры')
    doc.add_paragraph('   • {{passport_number}} - номер паспорта')
    doc.add_paragraph('   • {{article}} - артикул')
    doc.add_paragraph('   • {{matrix}} - матрица')
    doc.add_paragraph('   • И другие доступные поля')
    doc.add_paragraph('5. Сохраните шаблон')
    
    doc.add_paragraph()
    screenshot_para = doc.add_paragraph('[СКРИНШОТ: Редактор шаблона паспорта]')
    screenshot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    screenshot_para.runs[0].font.italic = True
    screenshot_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    
    doc.add_heading('8.2. Редактор шаблона наклейки', 2)
    doc.add_paragraph('1. Выберите вкладку "Шаблон наклейки"')
    doc.add_paragraph('2. Загрузите Excel или DOCX шаблон')
    doc.add_paragraph('3. Используйте плейсхолдеры для вставки данных')
    doc.add_paragraph('4. Настройте расположение элементов (логотип, штрихкоды)')
    doc.add_paragraph('5. Сохраните шаблон')
    
    doc.add_paragraph()
    screenshot_para = doc.add_paragraph('[СКРИНШОТ: Редактор шаблона наклейки]')
    screenshot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    screenshot_para.runs[0].font.italic = True
    screenshot_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    
    doc.add_paragraph()
    doc.add_paragraph('Примечание: Изменения в шаблонах применяются при следующей генерации документов.')
    
    doc.add_page_break()
    
    # Дополнительная информация
    doc.add_heading('Дополнительная информация', 1)
    
    doc.add_heading('Роли пользователей', 2)
    doc.add_paragraph('• Администратор (admin):')
    doc.add_paragraph('  - Полный доступ ко всем функциям')
    doc.add_paragraph('  - Управление пользователями и номенклатурой')
    doc.add_paragraph('  - Редактирование шаблонов')
    doc.add_paragraph('  - Просмотр всех паспортов в системе')
    
    doc.add_paragraph()
    doc.add_paragraph('• Пользователь (user):')
    doc.add_paragraph('  - Создание паспортов')
    doc.add_paragraph('  - Просмотр и экспорт своих паспортов')
    doc.add_paragraph('  - Создание и экспорт наклеек')
    
    doc.add_heading('Техническая поддержка', 2)
    doc.add_paragraph('При возникновении проблем или вопросов обращайтесь к администратору системы.')
    
    doc.add_heading('Безопасность', 2)
    doc.add_paragraph('• Не передавайте свои учетные данные третьим лицам')
    doc.add_paragraph('• Используйте надежные пароли')
    doc.add_paragraph('• Выходите из системы после завершения работы')
    doc.add_paragraph('• Сообщайте администратору о подозрительной активности')
    
    # Сохраняем документ
    output_path = 'user_manual_agb_passports.docx'
    doc.save(output_path)
    print(f"✅ Инструкция создана: {output_path}")
    return output_path

if __name__ == '__main__':
    create_user_manual()
