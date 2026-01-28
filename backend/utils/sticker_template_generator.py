"""
Генератор наклеек из шаблона Excel с плейсхолдерами
ПРОСТОЙ И ПРАВИЛЬНЫЙ АЛГОРИТМ:
1. Берем шаблон из backend/utils/templates/sticker_template.xlsx
2. Для каждой наклейки: заполняем шаблон данными + генерируем штрихкоды
    3. Конвертируем в DOCX и формируем таблицу 2x2 (2 колонки, 2 строки = 4 наклейки на странице)
"""
import os
import io
import tempfile
import sys
from typing import List, Optional

# Опциональный импорт docxtpl - если не установлен, будет использован fallback
try:
    from docxtpl import DocxTemplate, InlineImage
    DOCXTPL_AVAILABLE = True
except ImportError:
    DOCXTPL_AVAILABLE = False
    print("⚠️ docxtpl не установлен, будет использован стандартный метод генерации")

from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from backend.utils.template_manager import get_template_manager
import xml.etree.ElementTree as ET
from io import BytesIO
import shutil

# Импорт для работы с Excel
try:
    from openpyxl import load_workbook
    from openpyxl.drawing.image import Image as OpenpyxlImage
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    print("⚠️ openpyxl не установлен, будет использован стандартный метод генерации")

# Импорт Mm для использования в функциях
try:
    from docx.shared import Mm as DocxMm
except:
    DocxMm = Mm


def get_template_path():
    """Получает путь к шаблону наклеек через TemplateManager"""
    manager = get_template_manager()
    template_path = manager.get_template_path("sticker")
    if template_path:
        print(f"✅ Шаблон найден: {template_path}")
        return str(template_path)
    
    print("⚠️ Шаблон не найден, будет использован стандартный метод генерации")
    return None


def generate_stickers_from_template(passports, template_path=None):
    """
    Генерирует DOCX с наклейками из шаблона (Excel или DOCX)
    
    ПРОСТОЙ АЛГОРИТМ:
    1. Берем шаблон из templates/sticker_template.docx или .xlsx
    2. Для каждой наклейки заполняем шаблон данными + генерируем штрихкоды
    3. Формируем таблицу 2x2 (2 колонки, 2 строки = 4 наклейки на странице A4)
    """
    print(f"🏷️ Начинаем генерацию DOCX наклеек из шаблона для {len(passports)} паспортов")
    sys.stdout.flush()
    
    # Получаем путь к шаблону
    if template_path is None:
        template_path = get_template_path()
    
    if not template_path or not os.path.exists(template_path):
        print(f"⚠️ Шаблон не найден, используем стандартный метод")
        sys.stdout.flush()
        return generate_stickers_standard(passports)
    
    # Проверяем расширение файла
    if template_path.endswith('.xlsx'):
        if not OPENPYXL_AVAILABLE:
            print(f"⚠️ openpyxl не установлен, используем стандартный метод")
            sys.stdout.flush()
            return generate_stickers_standard(passports)
        
        try:
            result = generate_from_excel_template(passports, template_path)
            
            # Проверяем, что результат - это валидный DOCX (ZIP)
            import zipfile
            try:
                zip_buffer = io.BytesIO(result)
                with zipfile.ZipFile(zip_buffer, 'r') as zip_check:
                    if 'word/document.xml' not in zip_check.namelist():
                        raise ValueError("Результат не является валидным DOCX")
                print(f"✅ DOCX успешно сгенерирован из Excel шаблона, размер: {len(result)} байт")
                sys.stdout.flush()
                return result
            except (zipfile.BadZipFile, ValueError) as zip_err:
                print(f"❌ Результат не является валидным DOCX: {zip_err}")
                sys.stdout.flush()
                raise ValueError(f"Сгенерированный файл не является валидным DOCX: {zip_err}")
        except Exception as e:
            print(f"⚠️ Ошибка при использовании Excel шаблона: {e}")
            import traceback
            traceback.print_exc()
            sys.stdout.flush()
            print("🔄 Переключаемся на стандартный метод генерации DOCX...")
            sys.stdout.flush()
            return generate_stickers_standard(passports)
    
    # Fallback для DOCX шаблонов (старая логика)
    elif template_path.endswith('.docx'):
        if not DOCXTPL_AVAILABLE:
            print(f"⚠️ docxtpl не установлен, используем стандартный метод")
            sys.stdout.flush()
            return generate_stickers_standard(passports)
        
        try:
            result = generate_from_template_file(passports, template_path)
            
            # Проверяем, что результат - это валидный DOCX (ZIP)
            import zipfile
            try:
                zip_buffer = io.BytesIO(result)
                with zipfile.ZipFile(zip_buffer, 'r') as zip_check:
                    if 'word/document.xml' not in zip_check.namelist():
                        raise ValueError("Результат не является валидным DOCX")
                print(f"✅ DOCX успешно сгенерирован, размер: {len(result)} байт")
                sys.stdout.flush()
                return result
            except (zipfile.BadZipFile, ValueError) as zip_err:
                print(f"❌ Результат не является валидным DOCX: {zip_err}")
                sys.stdout.flush()
                raise ValueError(f"Сгенерированный файл не является валидным DOCX: {zip_err}")
        except Exception as e:
            print(f"⚠️ Ошибка при использовании шаблона: {e}")
            import traceback
            traceback.print_exc()
            sys.stdout.flush()
            print("🔄 Переключаемся на стандартный метод генерации DOCX...")
            sys.stdout.flush()
            return generate_stickers_standard(passports)
    else:
        print(f"⚠️ Неподдерживаемый формат шаблона: {template_path}")
        sys.stdout.flush()
        return generate_stickers_standard(passports)


def generate_from_excel_template(passports, template_path):
    """
    Генерирует DOCX наклейки из Excel шаблона
    
    АЛГОРИТМ:
    1. Загружаем Excel шаблон через openpyxl
    2. Для каждой наклейки заполняем ячейки данными
    3. Вставляем логотип и штрихкоды как изображения
    4. Конвертируем Excel в DOCX (через промежуточный формат)
    5. Формируем таблицу 2x4 в DOCX
    """
    print(f"📊 Генерация наклеек из Excel шаблона: {len(passports)} паспортов")
    sys.stdout.flush()
    
    from backend.utils.barcode_generator import generate_barcode_image
    
    # Получаем путь к логотипу
    manager = get_template_manager()
    logo_path_obj = manager.get_logo_path()
    logo_path = str(logo_path_obj) if logo_path_obj and logo_path_obj.exists() else None
    
    # Дополнительная проверка и логирование
    if not logo_path:
        # Пробуем альтернативные пути
        alt_paths = [
            '/app/backend/utils/templates/logo.png',
            '/app/templates/logo.png',
            'backend/utils/templates/logo.png',
            'templates/logo.png'
        ]
        for alt_path in alt_paths:
            if os.path.exists(alt_path):
                logo_path = alt_path
                print(f"    ✅ Логотип найден по альтернативному пути: {logo_path}")
                break
        
        if not logo_path:
            print(f"    ⚠️ Логотип не найден ни по одному из путей!")
    else:
        print(f"    ✅ Логотип найден: {logo_path}")
    
    # Создаем новый DOCX документ
    doc = Document()
    
    # Устанавливаем размеры страницы A4 без отступов
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(0)
    section.right_margin = Mm(0)
    section.top_margin = Mm(0)
    section.bottom_margin = Mm(0)
    
    # Загружаем Excel шаблон
    wb = load_workbook(template_path)
    ws = wb.active
    
    # Группируем паспорта по 4 на страницу (2x2)
    for page_idx in range(0, len(passports), 4):
        passport_group = passports[page_idx:page_idx+4]
        
        # Создаем таблицу 2x2 для страницы (2 колонки, 2 строки) = 4 наклейки на A4
        table = doc.add_table(rows=2, cols=2)
        
        # Настраиваем размеры таблицы и ячеек
        # A4 = 210мм × 297мм
        # 4 наклейки: 2 колонки × 2 строки
        # Ширина наклейки = 210мм / 2 = 105мм
        # Высота наклейки = 297мм / 2 = 148.5мм
        sticker_width = Mm(105)  # Ширина наклейки (полная ширина колонки)
        sticker_height = Mm(148.5)  # Высота наклейки = 297/2 (для 2 строк)
        
        # Настраиваем таблицу с границами
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Устанавливаем ширину таблицы на весь лист (210мм)
        tblWidth = OxmlElement('w:tblW')
        tblWidth.set(qn('w:w'), str(int(210 * 56.7)))  # 210мм в twips (1мм = 56.7 twips)
        tblWidth.set(qn('w:type'), 'dxa')
        tblPr.append(tblWidth)
        
        # Устанавливаем границы таблицы (тонкие линии между наклейками)
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Тонкая линия (0.5pt)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # Черный цвет
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Убираем отступы таблицы
        tblCellMar = OxmlElement('w:tblCellMar')
        for margin_name in ['top', 'left', 'bottom', 'right']:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), '0')
            margin.set(qn('w:type'), 'dxa')
            tblCellMar.append(margin)
        tblPr.append(tblCellMar)
        
        # Настраиваем ячейки
        for row in table.rows:
            # Устанавливаем высоту строки
            tr = row._tr
            trPr = tr.trPr
            if trPr is None:
                trPr = OxmlElement('w:trPr')
                tr.insert(0, trPr)
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(int(sticker_height * 20)))  # В twips (1мм = 20 twips)
            trHeight.set(qn('w:hRule'), 'exact')
            trPr.append(trHeight)
            
            for cell in row.cells:
                # Устанавливаем ширину ячейки
                tcPr = cell._element.tcPr
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    cell._element.insert(0, tcPr)
                
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(sticker_width * 56.7)))  # В twips
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)
                
                # Убираем все отступы в ячейках для точного заполнения
                tcMar = OxmlElement('w:tcMar')
                for margin_name in ['top', 'left', 'bottom', 'right']:
                    margin = OxmlElement(f'w:{margin_name}')
                    margin.set(qn('w:w'), '0')
                    margin.set(qn('w:type'), 'dxa')
                    tcMar.append(margin)
                tcPr.append(tcMar)
                
                # Вертикальное выравнивание по верху
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'top')
                tcPr.append(vAlign)
        
        # Заполняем ячейки (2x2 = 4 наклейки на странице)
        for row_idx in range(2):
            for col_idx in range(2):
                idx = row_idx * 2 + col_idx
                cell = table.rows[row_idx].cells[col_idx]
                
                if idx < len(passport_group):
                    passport = passport_group[idx]
                    nomenclature = passport.nomenclature
                    
                    if not nomenclature:
                        continue
                    
                    # Получаем данные
                    production_date = "2025"
                    if passport.created_at:
                        production_date = passport.created_at.strftime("%Y")
                    
                    # Генерируем штрихкоды
                    stock_code = nomenclature.article or getattr(nomenclature, 'code_1c', None) or '3501040'
                    stock_code_barcode_path = generate_barcode_image(stock_code, width_mm=40, height_mm=10)
                    
                    serial_number = passport.passport_number or 'AGB0000125'
                    serial_number_barcode_path = generate_barcode_image(serial_number, width_mm=40, height_mm=10)
                    
                    # СОЗДАЕМ ТАБЛИЦУ ВНУТРИ ЯЧЕЙКИ С ГРАНИЦАМИ, КАК В EXCEL ШАБЛОНЕ
                    # Очищаем ячейку полностью - удаляем все содержимое
                    # Удаляем все параграфы
                    for para in cell.paragraphs[:]:
                        p_element = para._element
                        p_element.getparent().remove(p_element)
                    
                    # Удаляем все таблицы (если есть)
                    for tbl in cell.tables[:]:
                        tbl_element = tbl._tbl
                        tbl_element.getparent().remove(tbl_element)
                    
                    # Создаем новый пустой параграф
                    cell.add_paragraph()
                    
                    # Создаем таблицу 10 строк x 3 колонки с границами (как в Excel шаблоне)
                    # В Excel шаблоне: 10 строк, 3 колонки
                    # Колонка A (1) - логотип (объединенная A1-A10)
                    # Колонка B (2) - метки/текст
                    # Колонка C (3) - значения (где нужно)
                    # ВАЖНО: таблица создается в последнем параграфе ячейки
                    try:
                        inner_table = cell.add_table(rows=10, cols=3)
                        inner_table.style = 'Table Grid'  # Стиль с границами
                        print(f"    ✅ Внутренняя таблица 10x3 создана для паспорта {passport.passport_number}")
                    except Exception as table_err:
                        print(f"    ❌ Ошибка создания внутренней таблицы: {table_err}")
                        import traceback
                        traceback.print_exc()
                        # Пропускаем эту наклейку
                        continue
                    
                    # Настраиваем ширины колонок
                    # Колонка A (логотип) - узкая, колонки B и C - шире
                    logo_col_width = sticker_width * 0.15  # 15% для логотипа
                    data_col_width = sticker_width * 0.425  # 42.5% для каждой колонки данных
                    for row in inner_table.rows:
                        # Колонка A (логотип)
                        row.cells[0].width = logo_col_width
                        # Колонки B и C (данные)
                        row.cells[1].width = data_col_width
                        row.cells[2].width = data_col_width
                            
                        # Настраиваем отступы для всех ячеек в строке
                        for col_idx in range(3):
                            inner_cell = row.cells[col_idx]
                            inner_tcPr = inner_cell._element.tcPr
                            if inner_tcPr is None:
                                inner_tcPr = OxmlElement('w:tcPr')
                                inner_cell._element.insert(0, inner_tcPr)
                            
                            # Устанавливаем ширину ячейки
                            width_value = logo_col_width if col_idx == 0 else data_col_width
                            tcW = OxmlElement('w:tcW')
                            tcW.set(qn('w:w'), str(int(width_value * 20)))  # В twips
                            tcW.set(qn('w:type'), 'dxa')
                            inner_tcPr.append(tcW)
                            
                            # Добавляем отступ 10 пикселей (примерно 3.75 мм) по периметру
                            # 10 пикселей при 96 DPI = примерно 3.75 мм = 75 twips (1 мм = 20 twips)
                            inner_tcMar = OxmlElement('w:tcMar')
                            for margin_name in ['top', 'left', 'bottom', 'right']:
                                margin = OxmlElement(f'w:{margin_name}')
                                margin.set(qn('w:w'), '75')  # 10 пикселей = ~3.75 мм = 75 twips
                                margin.set(qn('w:type'), 'dxa')
                                inner_tcMar.append(margin)
                            inner_tcPr.append(inner_tcMar)
                    
                    # Функция для установки шрифта Exo2
                    def set_exo2_font(run):
                        """Устанавливает шрифт Exo2 для текста"""
                        try:
                            run.font.name = 'Exo2'
                            run.font.size = Pt(10)
                        except:
                            # Если шрифт не найден, используем стандартный
                            pass
                    
                    # Получаем данные для формирования текста
                    matrix_val = nomenclature.matrix or 'HQ'
                    height_val = str(nomenclature.height or getattr(nomenclature, 'drilling_depth', None) or '12')
                    drilling_depth = getattr(nomenclature, 'drilling_depth', None) or height_val
                    # Извлекаем диапазон из drilling_depth (например, "11-13" из "11-13" или "09-12")
                    if '-' in str(drilling_depth):
                        depth_range = str(drilling_depth)
                    else:
                        # Если нет диапазона, пытаемся извлечь из серийного номера или используем высоту
                        depth_range = height_val
                    
                    # ОБЪЕДИНЯЕМ ВСЕ ЯЧЕЙКИ В ЛЕВОЙ КОЛОНКЕ (A1-A10) ДЛЯ ЛОГОТИПА
                    # Логотип должен быть слева, повернут на 90 градусов и вытянут вдоль всей высоты наклейки
                    # Колонка A (индекс 0) - для логотипа
                    try:
                        logo_cell = inner_table.rows[0].cells[0]  # A1
                        for row_idx in range(1, 10):
                            logo_cell.merge(inner_table.rows[row_idx].cells[0])  # Объединяем A1-A10
                        print(f"    ✅ Ячейки A1-A10 (колонка 0) объединены для логотипа")
                    except Exception as merge_err:
                        print(f"    ⚠️ Ошибка объединения ячеек для логотипа: {merge_err}")
                        import traceback
                        traceback.print_exc()
                        # Продолжаем без объединения
                        logo_cell = inner_table.rows[0].cells[0]
                    
                    # Логотип в объединенной левой колонке (повернут на 90 градусов, вытянут по высоте)
                    p_logo = logo_cell.paragraphs[0]
                    p_logo.alignment = 1  # CENTER (по вертикали будет по центру)
                    if logo_path and os.path.exists(logo_path):
                        try:
                            # Поворачиваем изображение перед добавлением через PIL
                            from PIL import Image
                            
                            # Открываем изображение
                            img = Image.open(logo_path)
                            
                            # Поворачиваем на 90 градусов против часовой стрелки (или -90)
                            # Для вертикального логотипа поворачиваем на -90 градусов
                            rotated_img = img.rotate(-90, expand=True)
                            
                            # Сохраняем во временный файл
                            temp_logo_path = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                            rotated_img.save(temp_logo_path.name)
                            temp_logo_path.close()
                            
                            # Высота наклейки = 148.5 мм, с учетом отступов (по 3.75 мм сверху и снизу) = 141 мм
                            # Ширина логотипа = примерно 12-15 мм
                            logo_height = Mm(141)  # Высота во всю наклейку
                            logo_width = Mm(12)    # Ширина логотипа
                            
                            logo_run = p_logo.add_run()
                            # Добавляем повернутое изображение
                            logo_run.add_picture(temp_logo_path.name, width=logo_width, height=logo_height)
                            
                            # Удаляем временный файл
                            try:
                                os.unlink(temp_logo_path.name)
                            except:
                                pass
                            
                            print(f"    ✅ Логотип добавлен вертикально (повернут на 90°) в левую колонку: {logo_path}")
                        except Exception as e:
                            print(f"    ⚠️ Ошибка добавления логотипа: {e}")
                            import traceback
                            traceback.print_exc()
                            # Fallback: добавляем без поворота
                            try:
                                logo_run = p_logo.add_run()
                                logo_run.add_picture(logo_path, width=Mm(12), height=Mm(141))
                            except:
                                pass
                    else:
                        print(f"    ⚠️ Логотип не найден для наклейки! Путь: {logo_path}")
                    
                    # СТРОКА 1: Пустая (или можно добавить что-то в B1, C1)
                    # В Excel шаблоне строка 1 пустая
                    
                    # СТРОКА 2: Название номенклатуры (B2, объединенная B2-C2)
                    row2 = inner_table.rows[1]  # Строка 2 (индекс 1)
                    cell_b2 = row2.cells[1]  # Колонка B (индекс 1)
                    cell_c2 = row2.cells[2]  # Колонка C (индекс 2)
                    cell_b2.merge(cell_c2)  # Объединяем B2 и C2
                    p_a2 = cell_b2.paragraphs[0]
                    p_a2.alignment = 1  # CENTER
                    
                    # СТРОКА 2: Название номенклатуры с форматированием "ALFA" (жирный + перенос строки)
                    # Используем только правую колонку B, так как A уже объединена с логотипом
                    row2 = inner_table.rows[1]
                    cell_b2 = row2.cells[1]  # Только правая колонка B
                    p_a2 = cell_b2.paragraphs[0]
                    p_a2.alignment = 1  # CENTER
                    
                    # Получаем название номенклатуры
                    nom_name = nomenclature.name or 'Коронка импрегнированная'
                    
                    # Форматируем: слово "ALFA" жирным и добавляем перенос строки после него
                    if 'ALFA' in nom_name.upper() or 'ALFA' in nom_name:
                        # Разделяем на части до и после ALFA
                        parts = nom_name.split('ALFA', 1)
                        if len(parts) == 2:
                            # До ALFA
                            run_before = p_a2.add_run(parts[0])
                            set_exo2_font(run_before)
                            # ALFA жирным
                            run_alfa = p_a2.add_run('ALFA')
                            set_exo2_font(run_alfa)
                            run_alfa.bold = True
                            # Перенос строки
                            p_a2.add_run('\n')
                            # После ALFA
                            run_after = p_a2.add_run(parts[1])
                            set_exo2_font(run_after)
                        else:
                            # Если не нашли ALFA, просто выводим как есть
                            run_a2 = p_a2.add_run(nom_name)
                            set_exo2_font(run_a2)
                    else:
                        # Если ALFA нет в названии, выводим как есть
                        run_a2 = p_a2.add_run(nom_name)
                        set_exo2_font(run_a2)
                    
                    # СТРОКА 3: Артикул (B3 = "Артикул:", C3 = значение + штрихкод)
                    row3 = inner_table.rows[2]  # Строка 3 (индекс 2)
                    cell_b3 = row3.cells[1]  # Колонка B
                    cell_c3 = row3.cells[2]  # Колонка C
                    p_b3 = cell_b3.paragraphs[0]
                    p_b3.alignment = 0  # LEFT
                    run_b3_label = p_b3.add_run('Артикул:')
                    set_exo2_font(run_b3_label)
                    p_c3 = cell_c3.paragraphs[0]
                    p_c3.alignment = 0  # LEFT
                    run_c3_value = p_c3.add_run(stock_code)
                    set_exo2_font(run_c3_value)
                    # Добавляем штрихкод артикула
                    if stock_code_barcode_path and os.path.exists(stock_code_barcode_path):
                        try:
                            p_c3.add_run('\n')
                            barcode_run = p_c3.add_run()
                            barcode_run.add_picture(stock_code_barcode_path, width=Mm(40), height=Mm(10))
                            print(f"    ✅ Штрихкод артикула добавлен: {stock_code}")
                        except Exception as e:
                            print(f"    ⚠️ Ошибка добавления штрихкода артикула: {e}")
                    
                    # СТРОКА 4: Высота матрицы (B4 = "Высота матрицы:", C4 = значение)
                    row4 = inner_table.rows[3]  # Строка 4 (индекс 3)
                    cell_b4 = row4.cells[1]  # Колонка B
                    cell_c4 = row4.cells[2]  # Колонка C
                    p_b4 = cell_b4.paragraphs[0]
                    p_b4.alignment = 0  # LEFT
                    run_b4_label = p_b4.add_run('Высота матрицы:')
                    set_exo2_font(run_b4_label)
                    p_c4 = cell_c4.paragraphs[0]
                    p_c4.alignment = 0  # LEFT
                    run_c4_value = p_c4.add_run(f"{height_val} мм.")
                    set_exo2_font(run_c4_value)
                    
                    # СТРОКА 5: Промывочные отверстия (B5 = "Промывочные отверстия:", C5 = значение)
                    row5 = inner_table.rows[4]  # Строка 5 (индекс 4)
                    cell_b5 = row5.cells[1]  # Колонка B
                    cell_c5 = row5.cells[2]  # Колонка C
                    p_b5 = cell_b5.paragraphs[0]
                    p_b5.alignment = 0  # LEFT
                    run_b5_label = p_b5.add_run('Промывочные отверстия:')
                    set_exo2_font(run_b5_label)
                    waterways_val = str(getattr(nomenclature, 'waterways', None) or '8')
                    p_c5 = cell_c5.paragraphs[0]
                    p_c5.alignment = 0  # LEFT
                    run_c5_value = p_c5.add_run(f"{waterways_val} мм.")
                    set_exo2_font(run_c5_value)
                    
                    # СТРОКА 6: Типоразмер (B6 = "Типоразмер:", C6 = значение)
                    row6 = inner_table.rows[5]  # Строка 6 (индекс 5)
                    cell_b6 = row6.cells[1]  # Колонка B
                    cell_c6 = row6.cells[2]  # Колонка C
                    p_b6 = cell_b6.paragraphs[0]
                    p_b6.alignment = 0  # LEFT
                    run_b6_label = p_b6.add_run('Типоразмер:')
                    set_exo2_font(run_b6_label)
                    p_c6 = cell_c6.paragraphs[0]
                    p_c6.alignment = 0  # LEFT
                    run_c6_value = p_c6.add_run(matrix_val)
                    set_exo2_font(run_c6_value)
                    
                    # СТРОКА 7: Серийный номер (B7 = "Серийный номер: значение", объединенная B7-C7)
                    row7 = inner_table.rows[6]  # Строка 7 (индекс 6)
                    cell_b7 = row7.cells[1]  # Колонка B
                    cell_c7 = row7.cells[2]  # Колонка C
                    cell_b7.merge(cell_c7)  # Объединяем B7 и C7
                    p_b7 = cell_b7.paragraphs[0]
                    p_b7.alignment = 0  # LEFT
                    run_b7_label = p_b7.add_run('Серийный номер: ')
                    set_exo2_font(run_b7_label)
                    run_b7_value = p_b7.add_run(serial_number)
                    set_exo2_font(run_b7_value)
                    if serial_number_barcode_path and os.path.exists(serial_number_barcode_path):
                        try:
                            p_b7.add_run('\n')
                            barcode_run = p_b7.add_run()
                            barcode_run.add_picture(serial_number_barcode_path, width=Mm(40), height=Mm(10))
                        except Exception as e:
                            print(f"    ⚠️ Ошибка добавления штрихкода serial_number: {e}")
                    
                    # СТРОКА 8: Дата изготовления (B8 = "Дата изготовления: значение", объединенная B8-C8)
                    row8 = inner_table.rows[7]  # Строка 8 (индекс 7)
                    cell_b8 = row8.cells[1]  # Колонка B
                    cell_c8 = row8.cells[2]  # Колонка C
                    cell_b8.merge(cell_c8)  # Объединяем B8 и C8
                    p_b8 = cell_b8.paragraphs[0]
                    p_b8.alignment = 0  # LEFT
                    # Формируем дату изготовления из даты создания паспорта
                    if passport.created_at:
                        day = passport.created_at.strftime("%d")
                        month = passport.created_at.strftime("%m")
                        year = passport.created_at.strftime("%Y")
                        date_text = f'Дата изготовления: «{day}» {month} {year}'
                    else:
                        from datetime import datetime
                        now = datetime.now()
                        day = now.strftime("%d")
                        month = now.strftime("%m")
                        year = now.strftime("%Y")
                        date_text = f'Дата изготовления: «{day}» {month} {year}'
                    run_b8 = p_b8.add_run(date_text)
                    set_exo2_font(run_b8)
                    
                    # СТРОКА 9: almazgeobur.ru (B9, объединенная B9-C9, по центру)
                    row9 = inner_table.rows[8]  # Строка 9 (индекс 8)
                    cell_b9 = row9.cells[1]  # Колонка B
                    cell_c9 = row9.cells[2]  # Колонка C
                    cell_b9.merge(cell_c9)  # Объединяем B9 и C9
                    p_b9 = cell_b9.paragraphs[0]
                    p_b9.alignment = 1  # CENTER
                    run_b9_site = p_b9.add_run('almazgeobur.ru')
                    set_exo2_font(run_b9_site)
                    
                    # СТРОКА 10: Пустая (как в Excel шаблоне)
                    
                    # Удаляем временные файлы штрихкодов
                    try:
                        if stock_code_barcode_path and os.path.exists(stock_code_barcode_path):
                            os.unlink(stock_code_barcode_path)
                        if serial_number_barcode_path and os.path.exists(serial_number_barcode_path):
                            os.unlink(serial_number_barcode_path)
                    except:
                        pass
        
        # Добавляем разрыв страницы (кроме последней)
        if page_idx + 4 < len(passports):
            doc.add_page_break()
    
    # Сохраняем в память
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    docx_content = buffer.getvalue()
    
    print(f"✅ DOCX с наклейками успешно сгенерирован из Excel, размер: {len(docx_content)} байт")
    sys.stdout.flush()
    
    return docx_content


def generate_from_template_file(passports, template_path):
    """
    ПРОСТОЙ И ПРАВИЛЬНЫЙ АЛГОРИТМ:
    1. Берем шаблон из backend/utils/templates/sticker_template.docx
    2. Для каждой наклейки:
       - Создаем копию шаблона
       - Генерируем штрихкоды
       - Заполняем через docxtpl
       - Копируем содержимое в ячейку таблицы 2x2
    3. Формируем таблицу 2x2 (2 колонки, 2 строки = 4 наклейки на странице)
    """
    print(f"🔍 Генерация наклеек: {len(passports)} паспортов, шаблон: {template_path}")
    sys.stdout.flush()
    
    if not DOCXTPL_AVAILABLE:
        print(f"⚠️ docxtpl не установлен, используем стандартный метод")
        return generate_stickers_standard(passports)
    
    # Загружаем шаблон для определения размеров
    template_doc = Document(template_path)
    template_section = template_doc.sections[0]
    
    # Определяем размеры страницы шаблона
    template_page_width_mm = template_section.page_width.mm if hasattr(template_section.page_width, 'mm') else template_section.page_width / 36000
    template_page_height_mm = template_section.page_height.mm if hasattr(template_section.page_height, 'mm') else template_section.page_height / 36000
    
    # Размеры наклейки из шаблона
    sticker_width_mm = template_page_width_mm
    sticker_height_mm = template_page_height_mm
    
    # Получаем путь к логотипу
    manager = get_template_manager()
    logo_path_obj = manager.get_logo_path()
    logo_path = str(logo_path_obj) if logo_path_obj and logo_path_obj.exists() else None
    
    # Дополнительная проверка и логирование
    if not logo_path:
        # Пробуем альтернативные пути
        alt_paths = [
            '/app/backend/utils/templates/logo.png',
            '/app/templates/logo.png',
            'backend/utils/templates/logo.png',
            'templates/logo.png'
        ]
        for alt_path in alt_paths:
            if os.path.exists(alt_path):
                logo_path = alt_path
                print(f"    ✅ Логотип найден по альтернативному пути: {logo_path}")
                break
        
        if not logo_path:
            print(f"    ⚠️ Логотип не найден ни по одному из путей!")
    else:
        print(f"    ✅ Логотип найден: {logo_path}")
    
    # Создаем новый документ
    doc = Document()
    
    # Устанавливаем размеры страницы A4 без отступов
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(0)
    section.right_margin = Mm(0)
    section.top_margin = Mm(0)
    section.bottom_margin = Mm(0)
    
    # Группируем паспорта по 4 на страницу (таблица 2x2)
    for page_idx in range(0, len(passports), 4):
        passport_group = passports[page_idx:page_idx+4]
        
        # Создаем таблицу 2x2 (2 строки, 2 колонки) = 4 наклейки на странице
        table = doc.add_table(rows=2, cols=2)
        table.style = None
        
        # Настраиваем таблицу
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Устанавливаем ширину таблицы (210мм = полная ширина страницы A4)
        total_table_width = 210  # Полная ширина страницы A4
        tblWidth = OxmlElement('w:tblW')
        tblWidth.set(qn('w:w'), str(int(total_table_width * 56.7)))
        tblWidth.set(qn('w:type'), 'dxa')
        tblPr.append(tblWidth)
        
        # Устанавливаем ширину колонок (по 105мм каждая для 2 колонок)
        tblGrid = OxmlElement('w:tblGrid')
        for col_idx in range(2):
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), str(int(105 * 56.7)))  # 105мм на колонку
            tblGrid.append(gridCol)
        tbl.append(tblGrid)
        
        # Устанавливаем границы таблицы для видимости (тонкие линии)
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Тонкая линия (0.5pt)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # Черный цвет
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Убираем отступы таблицы
        tblCellMar = OxmlElement('w:tblCellMar')
        for margin_name in ['top', 'left', 'bottom', 'right']:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), '0')
            margin.set(qn('w:type'), 'dxa')
            tblCellMar.append(margin)
        tblPr.append(tblCellMar)
        
        # Настраиваем ячейки (2x2 = 4 наклейки)
        # Высота наклейки = 297/2 = 148.5мм
        sticker_height_mm = 148.5
        sticker_width_mm = 105  # Ширина наклейки (полная ширина колонки)
        
        for row_idx in range(2):
            for col_idx in range(2):
                cell = table.rows[row_idx].cells[col_idx]
                tcPr = cell._element.tcPr
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    cell._element.insert(0, tcPr)
                
                # Устанавливаем ширину ячейки (105мм = полная ширина колонки)
                tcW = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcW')
                if tcW is None:
                    tcW = OxmlElement('w:tcW')
                    tcPr.append(tcW)
                tcW.set(qn('w:w'), str(int(sticker_width_mm * 56.7)))
                tcW.set(qn('w:type'), 'dxa')
                
                # Убираем отступы
                tcMar = OxmlElement('w:tcMar')
                for margin_name in ['top', 'left', 'bottom', 'right']:
                    margin = OxmlElement(f'w:{margin_name}')
                    margin.set(qn('w:w'), '0')
                    margin.set(qn('w:type'), 'dxa')
                    tcMar.append(margin)
                tcPr.append(tcMar)
                
                # Устанавливаем высоту строки (99мм)
                tr = cell._element.getparent()
                trPr = tr.trPr
                if trPr is None:
                    trPr = OxmlElement('w:trPr')
                    tr.insert(0, trPr)
                trHeight = OxmlElement('w:trHeight')
                trHeight.set(qn('w:val'), str(int(sticker_height_mm * 20)))
                trHeight.set(qn('w:hRule'), 'exact')
                trPr.append(trHeight)
                
                # Вертикальное выравнивание
                vAlign = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vAlign')
                if vAlign is None:
                    vAlign = OxmlElement('w:vAlign')
                    tcPr.append(vAlign)
                vAlign.set(qn('w:val'), 'top')
        
        # Заполняем таблицу наклейками
        print(f"📄 Страница {page_idx // 4 + 1}: заполняем таблицу 2x2 (4 наклейки)")
        sys.stdout.flush()
        
        for row_idx in range(2):
            for col_idx in range(2):
                idx = row_idx * 2 + col_idx
                cell = table.rows[row_idx].cells[col_idx]
                
                if idx < len(passport_group):
                    passport = passport_group[idx]
                    nomenclature = passport.nomenclature
                    
                    if not nomenclature:
                        print(f"⚠️ Нет номенклатуры для паспорта {passport.passport_number}")
                        continue
                    
                    # Получаем дату производства
                    production_date = "2025"
                    if passport.created_at:
                        production_date = passport.created_at.strftime("%Y")
                    
                    print(f"  📋 Наклейка [{row_idx}][{col_idx}]: {passport.passport_number}")
                    sys.stdout.flush()
                    
                    try:
                        # ПРОСТОЙ АЛГОРИТМ: рендерим шаблон и копируем содержимое
                        success = render_template_to_cell(
                            template_path,
                            cell,
                            doc,
                            passport,
                            nomenclature,
                            production_date,
                            logo_path
                        )
                        if success:
                            print(f"    ✅ Наклейка создана")
                        else:
                            print(f"    ⚠️ Ошибка создания наклейки")
                    except Exception as e:
                        print(f"    ❌ Ошибка: {e}")
                        import traceback
                        traceback.print_exc()
        
        # Добавляем разрыв страницы (кроме последней)
        if page_idx + 4 < len(passports):
            doc.add_page_break()
    
    # Сохраняем в память
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    docx_content = buffer.getvalue()
    
    print(f"✅ DOCX с наклейками успешно сгенерирован, размер: {len(docx_content)} байт")
    sys.stdout.flush()
    
    return docx_content


def render_template_to_cell(template_path, target_cell, target_doc, passport, nomenclature, production_date, logo_path):
    """
    ПРОСТОЙ И ПРАВИЛЬНЫЙ АЛГОРИТМ:
    1. Берем шаблон из backend/utils/templates/sticker_template.docx
    2. Генерируем штрихкоды для артикула и серийного номера
    3. Заполняем шаблон через docxtpl
    4. Копируем ВСЁ содержимое ячейки из рендеренного шаблона в целевую ячейку
    """
    try:
        from backend.utils.barcode_generator import generate_barcode_image
        import tempfile
        import shutil
        import zipfile
        import re
        
        print(f"    🔄 Рендерим шаблон для {passport.passport_number}...")
        
        # 1. Создаем копию шаблона
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            temp_template_path = tmp_file.name
            shutil.copy(template_path, temp_template_path)
        
        # 2. КРИТИЧНО: Исправляем проблемные плейсхолдеры ДО первого рендеринга
        # Это должно быть сделано ДО создания DocxTemplate
        try:
            with zipfile.ZipFile(temp_template_path, 'r') as zip_file:
                doc_xml = zip_file.read('word/document.xml')
                xml_str = doc_xml.decode('utf-8')
                
                # АГРЕССИВНОЕ исправление: удаляем ВСЕ проблемные конструкции
                fixed_xml = xml_str
                
                # 1. Удаляем .size из любых мест (включая внутри плейсхолдеров)
                fixed_xml = re.sub(r'\.size', '', fixed_xml)
                
                # 2. Удаляем проблемные фильтры Jinja2
                fixed_xml = re.sub(r'\|\s*\w+\([^)]*\)', '', fixed_xml)
                fixed_xml = re.sub(r'\|\s*\w+', '', fixed_xml)
                
                # 3. Исправляем разорванные плейсхолдеры - объединяем содержимое между {{ и }}
                # Ищем паттерны вида {{...<w:t>текст</w:t>...}} и заменяем на {{...текст...}}
                # Но делаем это аккуратно, не ломая XML структуру
                # Просто удаляем XML-теги внутри плейсхолдеров, сохраняя текст
                for _ in range(3):  # Повторяем несколько раз для вложенных тегов
                    # Ищем {{...<w:t>текст</w:t>...}} и заменяем на {{...текст...}}
                    fixed_xml = re.sub(r'(\{\{[^<]*?)(<w:t[^>]*>)([^<]*?)(</w:t>)([^}]*?\}\})', r'\1\3\5', fixed_xml)
                    fixed_xml = re.sub(r'(\{\{[^<]*?)(<w:rPr[^>]*>)([^<]*?)(</w:rPr>)([^}]*?\}\})', r'\1\5', fixed_xml)
                
                if fixed_xml != xml_str:
                    fixed_zip_path = temp_template_path + '.fixed'
                    with zipfile.ZipFile(fixed_zip_path, 'w', zipfile.ZIP_DEFLATED) as new_zip:
                        for item in zip_file.infolist():
                            if item.filename == 'word/document.xml':
                                new_zip.writestr(item, fixed_xml)
                            else:
                                new_zip.writestr(item, zip_file.read(item.filename))
                    
                    shutil.move(fixed_zip_path, temp_template_path)
                    print(f"    ✅ Шаблон исправлен (удалены проблемные конструкции)")
        except Exception as fix_err:
            print(f"    ⚠️ Ошибка исправления шаблона: {fix_err}")
            import traceback
            traceback.print_exc()
        
        # 3. Генерируем штрихкоды
        stock_code = nomenclature.article or getattr(nomenclature, 'code_1c', None) or '3501040'
        stock_code_barcode_path = generate_barcode_image(stock_code, width_mm=40, height_mm=10)
        
        serial_number = passport.passport_number or 'AGB0000125'
        serial_number_barcode_path = generate_barcode_image(serial_number, width_mm=40, height_mm=10)
        
        print(f"    📷 Штрихкоды сгенерированы: {stock_code}, {serial_number}")
        
        # 4. Загружаем шаблон ПОСЛЕ исправления
        template = DocxTemplate(temp_template_path)
        
        # Получаем данные из номенклатуры и паспорта
        # Формируем дату изготовления (печатная, не пустое поле)
        if passport.created_at:
            manufacture_date = passport.created_at.strftime("%d.%m.%Y")
        else:
            from datetime import datetime
            manufacture_date = datetime.now().strftime("%d.%m.%Y")
        
        matrix_val = nomenclature.matrix or 'HQ'
        height_val = str(nomenclature.height or getattr(nomenclature, 'drilling_depth', None) or '12')
        waterways_val = str(getattr(nomenclature, 'waterways', None) or '8')
        drilling_depth = getattr(nomenclature, 'drilling_depth', None) or height_val
        if '-' in str(drilling_depth):
            depth_range = str(drilling_depth)
        else:
            depth_range = height_val
        
        # Формируем название номенклатуры с форматированием ALFA (жирный + перенос строки)
        nom_name = nomenclature.name or 'Коронка импрегнированная'
        
        # Используем RichText для форматирования ALFA
        try:
            from docxtpl import RichText
            # Разделяем на части до и после ALFA
            if 'ALFA' in nom_name.upper() or 'ALFA' in nom_name:
                parts = nom_name.split('ALFA', 1)
                if len(parts) == 2:
                    nomenclature_name_rich = RichText()
                    nomenclature_name_rich.add(parts[0], font='Exo2')
                    nomenclature_name_rich.add('ALFA', font='Exo2', bold=True)
                    nomenclature_name_rich.add('\n', font='Exo2')  # Перенос строки
                    nomenclature_name_rich.add(parts[1], font='Exo2')
                    nomenclature_name = nomenclature_name_rich
                else:
                    nomenclature_name = nom_name
            else:
                nomenclature_name = nom_name
        except:
            # Если RichText не доступен, используем простой текст
            if 'ALFA' in nom_name.upper() or 'ALFA' in nom_name:
                parts = nom_name.split('ALFA', 1)
                if len(parts) == 2:
                    nomenclature_name = f"{parts[0]}ALFA\n{parts[1]}"
                else:
                    nomenclature_name = nom_name
            else:
                nomenclature_name = nom_name
        
        context = {
            'nomenclature_name': nomenclature_name,  # С форматированием ALFA
            'article': nomenclature.article or getattr(nomenclature, 'code_1c', None) or '3501040',
            'matrix': matrix_val,
            'height_mm': f"{height_val}ММ",  # Высота Матрица, мм: 12ММ
            'waterways_mm': f"{waterways_val}ММ",  # Промывочные отверстия; мм: 8ММ
            'serial_number': passport.passport_number or 'AGB 3-5 NQ 0000125',
            'serial number': passport.passport_number or 'AGB 3-5 NQ 0000125',  # С пробелом
            'waterways': waterways_val,
            'height': height_val,
            'production_date': production_date,
            'date': production_date,
            'manufacture_date': manufacture_date,  # Печатная дата изготовления
            'manufacture_date_formatted': f"{manufacture_date}г.",  # Дата изготовления: ДД.ММ.ГГГГг.
            'website': 'almazgeobur.ru',  # Без "Www"
            'order_number': getattr(passport, 'order_number', None) or '',
        }
        
        print(f"    📋 Данные для шаблона:")
        print(f"      - nomenclature_name: {context['nomenclature_name']}")
        print(f"      - article: {context['article']}")
        print(f"      - matrix: {context['matrix']}")
        print(f"      - serial_number: {context['serial_number']}")
        print(f"      - height: {context['height']}")
        print(f"      - waterways: {context['waterways']}")
        
        # 5. Добавляем изображения в контекст
        if logo_path and os.path.exists(logo_path):
            try:
                context['logo'] = InlineImage(template, logo_path, width=DocxMm(18), height=DocxMm(5.4))
                print(f"    ✅ Логотип добавлен в контекст: {logo_path}")
            except Exception as e:
                print(f"    ⚠️ Ошибка добавления логотипа: {e}")
                import traceback
                traceback.print_exc()
                context['logo'] = None
        else:
            print(f"    ⚠️ Логотип не найден: {logo_path}")
            context['logo'] = None
        
        if stock_code_barcode_path and os.path.exists(stock_code_barcode_path):
            try:
                context['stock_code'] = InlineImage(template, stock_code_barcode_path, width=DocxMm(40), height=DocxMm(10))
                print(f"    ✅ Штрихкод stock_code добавлен в контекст: {stock_code}")
            except Exception as e:
                print(f"    ⚠️ Ошибка добавления штрихкода stock_code: {e}")
                import traceback
                traceback.print_exc()
                context['stock_code'] = stock_code
        else:
            print(f"    ⚠️ Штрихкод stock_code не сгенерирован, используем текст: {stock_code}")
            context['stock_code'] = stock_code
        
        if serial_number_barcode_path and os.path.exists(serial_number_barcode_path):
            try:
                context['serial_number_code'] = InlineImage(template, serial_number_barcode_path, width=DocxMm(40), height=DocxMm(10))
                print(f"    ✅ Штрихкод serial_number_code добавлен в контекст: {serial_number}")
            except Exception as e:
                print(f"    ⚠️ Ошибка добавления штрихкода serial_number_code: {e}")
                import traceback
                traceback.print_exc()
                context['serial_number_code'] = serial_number
        else:
            print(f"    ⚠️ Штрихкод serial_number_code не сгенерирован, используем текст: {serial_number}")
            context['serial_number_code'] = serial_number
        
        # 6. Рендерим шаблон
        render_success = False
        try:
            template.render(context)
            print(f"    ✅ Шаблон отрендерен успешно")
            render_success = True
        except Exception as render_err:
            error_str = str(render_err)
            print(f"    ⚠️ Ошибка рендеринга: {error_str}")
            import traceback
            traceback.print_exc()
            
            # Пробуем еще раз с более агрессивным исправлением
            try:
                print(f"    🔄 Пробуем более агрессивное исправление шаблона...")
                with zipfile.ZipFile(temp_template_path, 'r') as zip_file:
                    doc_xml = zip_file.read('word/document.xml')
                    xml_str = doc_xml.decode('utf-8')
                    
                    # Еще более агрессивное исправление
                    fixed_xml = xml_str
                    # Удаляем все .size
                    fixed_xml = re.sub(r'\.size', '', fixed_xml)
                    # Удаляем все фильтры
                    fixed_xml = re.sub(r'\|\s*\w+\([^)]*\)', '', fixed_xml)
                    fixed_xml = re.sub(r'\|\s*\w+', '', fixed_xml)
                    
                    if fixed_xml != xml_str:
                        fixed_zip_path = temp_template_path + '.fixed2'
                        with zipfile.ZipFile(fixed_zip_path, 'w', zipfile.ZIP_DEFLATED) as new_zip:
                            for item in zip_file.infolist():
                                if item.filename == 'word/document.xml':
                                    new_zip.writestr(item, fixed_xml)
                                else:
                                    new_zip.writestr(item, zip_file.read(item.filename))
                        shutil.move(fixed_zip_path, temp_template_path)
                        template = DocxTemplate(temp_template_path)
                        # Обновляем изображения в контексте
                        if logo_path and os.path.exists(logo_path):
                            context['logo'] = InlineImage(template, logo_path, width=DocxMm(18), height=DocxMm(5.4))
                        if stock_code_barcode_path and os.path.exists(stock_code_barcode_path):
                            context['stock_code'] = InlineImage(template, stock_code_barcode_path, width=DocxMm(40), height=DocxMm(10))
                        if serial_number_barcode_path and os.path.exists(serial_number_barcode_path):
                            context['serial_number_code'] = InlineImage(template, serial_number_barcode_path, width=DocxMm(40), height=DocxMm(10))
                        
                        # Пробуем рендерить снова
                        try:
                            template.render(context)
                            print(f"    ✅ Шаблон отрендерен после исправления")
                            render_success = True
                        except Exception as render_err2:
                            print(f"    ⚠️ Повторный рендеринг тоже не удался: {render_err2}")
                            # Продолжаем с неотрендеренным шаблоном
            except Exception as retry_err:
                print(f"    ⚠️ Повторная попытка исправления не удалась: {retry_err}")
                # Продолжаем с неотрендеренным шаблоном
        
        if not render_success:
            print(f"    ⚠️ ВНИМАНИЕ: Шаблон не был отрендерен, данные могут быть не заполнены!")
            # Если рендеринг не работает, заполняем данные вручную в XML
            try:
                print(f"    🔄 Пробуем заполнить данные вручную в XML...")
                with zipfile.ZipFile(temp_template_path, 'r') as zip_file:
                    doc_xml = zip_file.read('word/document.xml')
                    xml_str = doc_xml.decode('utf-8')
                    
                    # Заменяем плейсхолдеры на данные
                    replacements = {
                        '{{nomenclature_name}}': context.get('nomenclature_name', 'Коронка импрегнированная HQ 11-13 высота 12 мм'),
                        '{{article}}': context.get('article', ''),
                        '{{matrix}}': context.get('matrix', 'HQ'),
                        '{{height_mm}}': context.get('height_mm', '12ММ'),
                        '{{waterways_mm}}': context.get('waterways_mm', '8ММ'),
                        '{{serial_number}}': context.get('serial_number', ''),
                        '{{serial number}}': context.get('serial number', ''),
                        '{{waterways}}': context.get('waterways', '8'),
                        '{{height}}': context.get('height', '12'),
                        '{{production_date}}': context.get('production_date', ''),
                        '{{date}}': context.get('date', ''),
                        '{{manufacture_date}}': context.get('manufacture_date', ''),
                        '{{manufacture_date_formatted}}': context.get('manufacture_date_formatted', ''),
                        '{{website}}': context.get('website', 'almazgeobur.ru'),
                    }
                    
                    for placeholder, value in replacements.items():
                        xml_str = xml_str.replace(placeholder, str(value))
                    
                    # Сохраняем исправленный XML
                    fixed_zip_path = temp_template_path + '.manual_fixed'
                    with zipfile.ZipFile(fixed_zip_path, 'w', zipfile.ZIP_DEFLATED) as new_zip:
                        for item in zip_file.infolist():
                            if item.filename == 'word/document.xml':
                                new_zip.writestr(item, xml_str)
                            else:
                                new_zip.writestr(item, zip_file.read(item.filename))
                    shutil.move(fixed_zip_path, temp_template_path)
                    template = DocxTemplate(temp_template_path)
                    print(f"    ✅ Данные заполнены вручную в XML")
            except Exception as manual_err:
                print(f"    ⚠️ Ошибка ручного заполнения данных: {manual_err}")
        
        # 7. Сохраняем рендеренный шаблон
        rendered_path = temp_template_path.replace('.docx', '_rendered.docx')
        template.save(rendered_path)
        
        # 8. Загружаем рендеренный документ
        try:
            rendered_doc = Document(rendered_path)
        except Exception as load_err:
            print(f"    ❌ Ошибка загрузки рендеренного документа: {load_err}")
            import traceback
            traceback.print_exc()
            return False
        
        if not rendered_doc.tables or len(rendered_doc.tables) == 0:
            print(f"    ❌ Рендеренный шаблон не содержит таблицы")
            return False
        
        rendered_table = rendered_doc.tables[0]
        if len(rendered_table.rows) == 0 or len(rendered_table.rows[0].cells) == 0:
            print(f"    ❌ Рендеренный шаблон не содержит ячейки")
            return False
        
        source_cell = rendered_table.rows[0].cells[0]
        
        # Убеждаемся, что это объекты, а не строки
        if not hasattr(rendered_doc, 'part'):
            print(f"    ❌ rendered_doc не имеет атрибута part")
            return False
        source_part = rendered_doc.part
        
        if not hasattr(target_doc, 'part'):
            print(f"    ❌ target_doc не имеет атрибута part")
            return False
        target_part = target_doc.part
        
        # Проверяем, что это правильные объекты
        if not hasattr(source_part, 'rels'):
            print(f"    ❌ source_part не имеет атрибута rels: {type(source_part)}")
            return False
        if not hasattr(target_part, 'rels'):
            print(f"    ❌ target_part не имеет атрибута rels: {type(target_part)}")
            return False
        
        # 9. Копируем ВСЁ содержимое ячейки из рендеренного шаблона
        # Очищаем целевую ячейку
        for child in list(target_cell._element):
            if not child.tag.endswith('}tcPr'):
                target_cell._element.remove(child)
        
        # Собираем все элементы ячейки (кроме tcPr)
        source_content = []
        for child in source_cell._element:
            if not child.tag.endswith('}tcPr'):
                element_xml_str = ET.tostring(child, encoding='unicode')
                element_xml = ET.fromstring(element_xml_str)
                source_content.append((child.tag, element_xml))
        
        print(f"    📋 Найдено {len(source_content)} элементов для копирования")
        
        # Вставляем все элементы в целевую ячейку с обработкой изображений
        images_copied = 0
        
        # Сначала обрабатываем все изображения во всех элементах
        for tag, element_xml in source_content:
            all_drawings = element_xml.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
            if len(all_drawings) > 0:
                print(f"    🔍 В элементе {tag.split('}')[-1]}: найдено {len(all_drawings)} изображений")
            
            for drawing in all_drawings:
                blip = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                if blip is not None:
                    embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed_id:
                        # Проверяем relationships напрямую
                        if embed_id in source_part.rels:
                            try:
                                image_rel = source_part.rels[embed_id]
                                # Проверяем, что это действительно изображение
                                if hasattr(image_rel, 'target_part') and hasattr(image_rel.target_part, 'blob'):
                                    image_blob = image_rel.target_part.blob
                                    # Добавляем изображение в целевой документ через правильный API
                                    try:
                                        # relate_to требует файл, а не BytesIO - создаем временный файл
                                        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_img:
                                            tmp_img.write(image_blob)
                                            tmp_img_path = tmp_img.name
                                        
                                        # Используем relate_to с файлом
                                        rId = target_part.relate_to(tmp_img_path, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image', is_external=False)
                                        new_embed_id = rId
                                        
                                        # Удаляем временный файл
                                        try:
                                            os.unlink(tmp_img_path)
                                        except:
                                            pass
                                    except Exception as add_img_err:
                                        print(f"    ⚠️ Ошибка добавления изображения: {add_img_err}")
                                        import traceback
                                        traceback.print_exc()
                                        continue
                                    # Обновляем embed_id в XML
                                    blip.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed', new_embed_id)
                                    images_copied += 1
                                    print(f"    ✅ Изображение {images_copied} скопировано (embed_id: {embed_id} -> {new_embed_id})")
                                else:
                                    print(f"    ⚠️ Relationship {embed_id} не содержит изображение")
                            except Exception as img_err:
                                print(f"    ⚠️ Ошибка копирования изображения {embed_id}: {img_err}")
                                import traceback
                                traceback.print_exc()
                        else:
                            # Пробуем найти изображение по содержимому relationships
                            found = False
                            if not hasattr(source_part, 'rels'):
                                print(f"    ⚠️ source_part не имеет rels: {type(source_part)}")
                                continue
                            for rel_id, rel in source_part.rels.items():
                                try:
                                    if hasattr(rel, 'target_part') and hasattr(rel.target_part, 'content_type'):
                                        if 'image' in rel.target_part.content_type or 'png' in rel.target_part.content_type or 'jpeg' in rel.target_part.content_type:
                                            # Проверяем размер - если похож на наше изображение
                                            if hasattr(rel.target_part, 'blob'):
                                                image_blob = rel.target_part.blob
                                                if len(image_blob) > 100:  # Минимальный размер изображения
                                                    try:
                                                        # relate_to требует файл, а не BytesIO - создаем временный файл
                                                        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_img:
                                                            tmp_img.write(image_blob)
                                                            tmp_img_path = tmp_img.name
                                                        
                                                        # Используем relate_to с файлом
                                                        rId = target_part.relate_to(tmp_img_path, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image', is_external=False)
                                                        new_embed_id = rId
                                                        
                                                        # Удаляем временный файл
                                                        try:
                                                            os.unlink(tmp_img_path)
                                                        except:
                                                            pass
                                                    except Exception as add_img_err:
                                                        print(f"    ⚠️ Ошибка добавления изображения: {add_img_err}")
                                                        continue
                                                    blip.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed', new_embed_id)
                                                    images_copied += 1
                                                    print(f"    ✅ Изображение {images_copied} найдено и скопировано (rel_id: {rel_id})")
                                                    found = True
                                                    break
                                except:
                                    pass
                            
                            if not found:
                                print(f"    ⚠️ embed_id {embed_id} не найден в relationships")
        
        # Теперь копируем все элементы с уже обработанными изображениями
        for tag, element_xml in source_content:
            try:
                element_xml_str = ET.tostring(element_xml, encoding='unicode')
                copied_element = parse_xml(element_xml_str)
                target_cell._element.append(copied_element)
                element_name = tag.split('}')[-1] if '}' in tag else tag
                print(f"    ✅ Элемент {element_name} скопирован")
            except Exception as append_err:
                print(f"    ⚠️ Ошибка копирования элемента {tag}: {append_err}")
                import traceback
                traceback.print_exc()
        
        print(f"    ✅ Содержимое скопировано: {len(source_content)} элементов, {images_copied} изображений")
        
        # Удаляем временные файлы
        try:
            os.unlink(temp_template_path)
            os.unlink(rendered_path)
            if stock_code_barcode_path and os.path.exists(stock_code_barcode_path):
                os.unlink(stock_code_barcode_path)
            if serial_number_barcode_path and os.path.exists(serial_number_barcode_path):
                os.unlink(serial_number_barcode_path)
        except:
            pass
        
        return True
        
    except Exception as e:
        print(f"    ❌ Ошибка при рендеринге шаблона: {e}")
        import traceback
        traceback.print_exc()
        return False


# Стандартный метод генерации (fallback)
def generate_stickers_standard(passports):
    """Стандартный метод генерации наклеек без шаблона"""
    print(f"🔄 Используем стандартный метод генерации для {len(passports)} паспортов")
    sys.stdout.flush()
    
    # Создаем простой документ
    doc = Document()
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    
    # Группируем по 4 на страницу (2x2)
    for page_idx in range(0, len(passports), 4):
        passport_group = passports[page_idx:page_idx+4]
        table = doc.add_table(rows=2, cols=2)
        
        for row_idx in range(2):
            for col_idx in range(2):
                idx = row_idx * 2 + col_idx
                cell = table.rows[row_idx].cells[col_idx]
                
                if idx < len(passport_group):
                    passport = passport_group[idx]
                    nomenclature = passport.nomenclature
                    if nomenclature:
                        # Форматируем название с ALFA жирным и переносом строки
                        nom_name = nomenclature.name or 'Коронка импрегнированная'
                        if 'ALFA' in nom_name.upper() or 'ALFA' in nom_name:
                            parts = nom_name.split('ALFA', 1)
                            if len(parts) == 2:
                                p = cell.paragraphs[0]
                                p.clear()
                                run_before = p.add_run(parts[0])
                                run_alfa = p.add_run('ALFA')
                                run_alfa.bold = True
                                p.add_run('\n')
                                run_after = p.add_run(parts[1])
                                p.add_run(f"\n{passport.passport_number}")
                            else:
                                cell.text = f"{nom_name}\n{passport.passport_number}"
                        else:
                            cell.text = f"{nom_name}\n{passport.passport_number}"
        
        if page_idx + 4 < len(passports):
            doc.add_page_break()
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
