import io
import os
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer, Image, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
# from cairosvg import svg2png  # Временно отключено для локальной разработки
from reportlab.lib.utils import ImageReader


def create_logo_image():
    """Создает изображение логотипа из PNG файла
    
    ВАЖНО: Приоритет отдается backend/utils/templates/logo.png
    """
    try:
        import os
        
        # Получаем текущую директорию (backend/utils/)
        current_dir = os.path.dirname(os.path.abspath(__file__))
        
        # ПРИОРИТЕТ 1: Логотип из /app/backend/utils/templates/logo.png (в Docker контейнере - АБСОЛЮТНЫЙ ПУТЬ)
        docker_templates_path = '/app/backend/utils/templates/logo.png'
        if os.path.exists(docker_templates_path):
            file_size = os.path.getsize(docker_templates_path)
            if file_size > 0:
                print(f"✅ Логотип найден в /app/backend/utils/templates/logo.png (размер: {file_size} байт)")
                return docker_templates_path
            else:
                print(f"⚠️ Файл логотипа пустой: {docker_templates_path}")
        
        # ПРИОРИТЕТ 2: Логотип из backend/utils/templates/logo.png (относительный путь)
        templates_logo_path = os.path.join(current_dir, 'templates', 'logo.png')
        if os.path.exists(templates_logo_path):
            file_size = os.path.getsize(templates_logo_path)
            if file_size > 0:
                print(f"✅ Логотип найден в backend/utils/templates/logo.png (размер: {file_size} байт)")
                return templates_logo_path
            else:
                print(f"⚠️ Файл логотипа пустой: {templates_logo_path}")
        
        # ПРИОРИТЕТ 3: Логотип из /app/templates/logo.png (fallback для старых версий)
        template_logo_path = '/app/templates/logo.png'
        if os.path.exists(template_logo_path):
            file_size = os.path.getsize(template_logo_path)
            if file_size > 0:
                print(f"✅ Логотип найден в /app/templates/logo.png (размер: {file_size} байт)")
                # Пытаемся скопировать в правильное место для будущего использования
                try:
                    import shutil
                    os.makedirs('/app/backend/utils/templates', exist_ok=True)
                    shutil.copy2(template_logo_path, docker_templates_path)
                    print(f"✅ Логотип скопирован в {docker_templates_path}")
                    return docker_templates_path
                except Exception as copy_err:
                    print(f"⚠️ Не удалось скопировать логотип: {copy_err}")
                    return template_logo_path
            else:
                print(f"⚠️ Файл логотипа пустой: {template_logo_path}")
        
        # ПРИОРИТЕТ 4: Логотип из backend/utils/logo.png
        utils_logo_path = os.path.join(current_dir, 'logo.png')
        if os.path.exists(utils_logo_path):
            file_size = os.path.getsize(utils_logo_path)
            if file_size > 0:
                print(f"✅ Логотип найден в backend/utils/logo.png (размер: {file_size} байт)")
                # Пытаемся скопировать в templates для будущего использования
                try:
                    import shutil
                    os.makedirs(os.path.join(current_dir, 'templates'), exist_ok=True)
                    shutil.copy2(utils_logo_path, templates_logo_path)
                    print(f"✅ Логотип скопирован в {templates_logo_path}")
                    return templates_logo_path
                except Exception as copy_err:
                    print(f"⚠️ Не удалось скопировать логотип: {copy_err}")
                    return utils_logo_path
        
        # Fallback: другие возможные пути
        project_root = os.path.dirname(os.path.dirname(current_dir))
        fallback_paths = [
            os.path.join(project_root, 'logo.png'),
            '/root/agb_passports/logo.png',
            '/root/agb_passports/backend/utils/logo.png',
            '/root/agb_passports/backend/utils/templates/logo.png',
        ]
        
        for logo_path in fallback_paths:
            abs_path = os.path.abspath(logo_path) if not logo_path.startswith('/') else logo_path
            if os.path.exists(abs_path):
                file_size = os.path.getsize(abs_path)
                if file_size > 0:
                    print(f"✅ Логотип найден (fallback): {abs_path} (размер: {file_size} байт)")
                    # Пытаемся скопировать в правильное место
                    try:
                        import shutil
                        # Копируем в backend/utils/templates/logo.png
                        os.makedirs(os.path.join(current_dir, 'templates'), exist_ok=True)
                        shutil.copy2(abs_path, templates_logo_path)
                        print(f"✅ Логотип скопирован в {templates_logo_path}")
                        return templates_logo_path
                    except Exception as copy_err:
                        print(f"⚠️ Не удалось скопировать логотип: {copy_err}")
                        return abs_path
        
        # Если файл не существует, пытаемся создать его
        print("⚠️ Логотип не найден, пытаемся создать...")
        try:
            from backend.utils.create_logo import create_logo_png
            created_path = create_logo_png()
            # Если создан в другом месте, копируем в правильное
            if created_path and created_path != templates_logo_path:
                try:
                    import shutil
                    os.makedirs(os.path.join(current_dir, 'templates'), exist_ok=True)
                    shutil.copy2(created_path, templates_logo_path)
                    print(f"✅ Логотип скопирован в {templates_logo_path}")
                    return templates_logo_path
                except:
                    return created_path
            return created_path
        except Exception as create_error:
            print(f"⚠️ Не удалось создать логотип: {create_error}")
        
        print("❌ Не удалось найти или создать логотип")
        return None
    except Exception as e:
        print(f"❌ Ошибка при загрузке логотипа: {e}")
        import traceback
        traceback.print_exc()
        return None


def setup_cyrillic_fonts():
    """Настраивает шрифты с поддержкой кириллицы"""
    try:
        # Пробуем разные пути к шрифтам
        font_paths = [
            '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
            '/usr/share/fonts/TTF/DejaVuSans.ttf',
            '/usr/share/fonts/dejavu/DejaVuSans.ttf',
            '/System/Library/Fonts/Arial.ttf',  # macOS
            '/System/Library/Fonts/Helvetica.ttc',  # macOS
            '/Windows/Fonts/arial.ttf',  # Windows
            '/Windows/Fonts/calibri.ttf',  # Windows
            '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',  # Linux
            '/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf',  # Linux
        ]
        
        normal_font = 'Helvetica'
        for path in font_paths:
            try:
                if os.path.exists(path):
                    pdfmetrics.registerFont(TTFont('DejaVuSans', path))
                    normal_font = 'DejaVuSans'
                    print(f"✅ Успешно загружен шрифт: {path}")
                    break
            except Exception as e:
                print(f"❌ Ошибка загрузки шрифта {path}: {e}")
                continue

        # Если DejaVu не найден, попробуем другие шрифты
        if normal_font == 'Helvetica':
            try:
                arial_paths = [
                    '/usr/share/fonts/truetype/msttcorefonts/Arial_Unicode_MS.ttf',
                    '/System/Library/Fonts/Arial.ttf',
                    '/Windows/Fonts/arial.ttf'
                ]
                for path in arial_paths:
                    if os.path.exists(path):
                        pdfmetrics.registerFont(TTFont('ArialUnicode', path))
                        normal_font = 'ArialUnicode'
                        print(f"✅ Успешно загружен Arial шрифт: {path}")
                        break
            except Exception as e:
                print(f"❌ Ошибка загрузки Arial шрифта: {e}")
                pass
        
        # Если ничего не найдено, используем встроенные шрифты ReportLab
        if normal_font == 'Helvetica':
            print("⚠️ Внешние шрифты не найдены, используем встроенные")
            # Регистрируем встроенные шрифты с правильной кодировкой
            try:
                from reportlab.pdfbase.cidfonts import UnicodeCIDFont
                pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
                normal_font = 'STSong-Light'
                print("✅ Используем встроенный Unicode шрифт")
            except:
                # Попробуем создать простой шрифт с поддержкой кириллицы
                try:
                    # Используем Times-Roman как fallback с правильной кодировкой
                    normal_font = 'Times-Roman'
                    print("✅ Используем Times-Roman с UTF-8 кодировкой")
                except:
                    print("❌ Не удалось загрузить Unicode шрифт, используем Helvetica")
        
        print(f"📝 Используемый шрифт: {normal_font}")
    except Exception as e:
        print(f"❌ Общая ошибка настройки шрифтов: {e}")
        normal_font = 'Helvetica'
    
    return normal_font


def create_passport_styles(normal_font):
    """Создает стили для PDF паспорта"""
    styles = getSampleStyleSheet()
    
    # Стили для заголовков и текста с корректной кодировкой
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=20,
        alignment=1,  # Центрирование
        fontName=normal_font,
        encoding='utf-8'
    )

    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=15,
        fontName=normal_font,
        encoding='utf-8'
    )

    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontName=normal_font,
        fontSize=10,
        encoding='utf-8'
    )
    
    return title_style, subtitle_style, normal_style


def create_passport_pdf_content(passport, normal_font, title_style, subtitle_style, normal_style):
    """Создает содержимое PDF паспорта в новом формате согласно образцу"""
    story = []

    # Создаем контактную информацию в одну строку
    contact_info = """ООО "Алмазгеобур" 125362, г. Москва, улица Водников, дом 2, стр. 14, оф. 11, тел.:+7 495 229 82 94
LLP "Almazgeobur" 125362, Moscow, Vodnikov Street, 2, building. 14, of. 11, tel.:+7 495 229 82 94,
e-mail: contact@almazgeobur.ru"""

    # Создаем внешнюю таблицу для общей рамки
    outer_data = [[None]]  # Будет заполнено позже
    outer_table = Table(outer_data, colWidths=[190*mm])
    outer_table.setStyle(TableStyle([
        ('BOX', (0, 0), (-1, -1), 0.5, colors.black),
        ('TOPPADDING', (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ('LEFTPADDING', (0, 0), (-1, -1), 2),
        ('RIGHTPADDING', (0, 0), (-1, -1), 2),
    ]))

    # Создаем внутреннюю таблицу с логотипом и контактными данными
    header_data = [[None, contact_info]]

    # Добавляем логотип (точно такой же размер, как в наклейках: 18мм × 5.4мм)
    logo_img = create_logo_image()
    logo_added = False
    if logo_img and os.path.exists(logo_img):
        try:
            print(f"✅ Логотип найден для паспорта: {logo_img}")
            print(f"   Размеры логотипа: 18мм × 5.4мм (как в наклейках)")
            # Проверяем, что файл действительно является изображением
            file_size = os.path.getsize(logo_img)
            print(f"   Размер файла: {file_size} байт")
            # Используем точно те же размеры, что и в наклейках
            logo_cell = Image(logo_img, width=18*mm, height=5.4*mm)
            print(f"   Логотип добавлен в паспорт с размерами: {18*mm} × {5.4*mm} точек")
            header_data[0][0] = logo_cell
            logo_added = True
            print(f"   ✅ Логотип успешно добавлен в header_data")
        except Exception as e:
            print(f"⚠️ Ошибка при добавлении логотипа в паспорт: {e}")
            import traceback
            traceback.print_exc()
            # Пытаемся использовать ImageReader для более надежной загрузки
            try:
                from reportlab.lib.utils import ImageReader
                logo_cell = Image(ImageReader(logo_img), width=18*mm, height=5.4*mm)
                header_data[0][0] = logo_cell
                logo_added = True
                print(f"   ✅ Логотип добавлен через ImageReader")
            except Exception as e2:
                print(f"⚠️ Ошибка при добавлении логотипа через ImageReader: {e2}")
    
    if not logo_added:
        print(f"⚠️ Логотип не найден для паспорта: {logo_img}")
        # Если логотип не найден, оставляем пустую ячейку (не добавляем текстовый fallback)
        # Это предотвратит отображение текстового логотипа
        header_data[0][0] = ""  # Пустая строка вместо None

    # Создаем заголовочную таблицу ПОСЛЕ добавления логотипа
        header_table = Table(header_data, colWidths=[45*mm, 143*mm])
        header_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), normal_font),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (0, 0), 'LEFT'),
            ('ALIGN', (1, 0), (1, 0), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('ENCODING', (0, 0), (-1, -1), 'utf-8'),
        ]))

    story.append(header_table)
    story.append(Spacer(1, 15))

    # Основная таблица с данными паспорта (как на картинке)
    # Получаем данные из БД
    nomenclature = passport.nomenclature
    if not nomenclature:
        print(f"❌ Номенклатура не найдена для паспорта {passport.passport_number}")
        return story
    
    # Генерируем штрихкод
    barcode = f"AGB{nomenclature.article or '3501040'}-{passport.passport_number or '0000125'}"

    # Создаем стиль для переноса текста
    wrapped_style = ParagraphStyle(
        'WrappedText',
        parent=normal_style,
        fontSize=7,  # Уменьшенный размер шрифта
        leading=9,   # Межстрочный интервал
        spaceBefore=0,
        spaceAfter=0,
    )

    # Используем наименование инструмента из номенклатуры
    tool_type_name = nomenclature.name or "Буровой инструмент / Drilling tool"
    
    # Создаем стиль для ячеек с переносом текста
    cell_style = ParagraphStyle(
        'CellText',
        parent=normal_style,
        fontSize=7,
        leading=9,
        spaceBefore=0,
        spaceAfter=0,
        alignment=1,  # CENTER
    )
    
    # Создаем параграф с названием номенклатуры для автоматического переноса
    tool_type = Paragraph(tool_type_name, cell_style)

    passport_data = [
        [Paragraph("Артикул / Stock Code", cell_style), 
         Paragraph("Типоразмер / Tool size", cell_style), 
         Paragraph("Серийный номер / Serial Number", cell_style), 
         Paragraph("Буровой инструмент / Tool type", cell_style)],
        [Paragraph(nomenclature.article or "3501040", cell_style), 
         Paragraph(nomenclature.matrix or "NQ", cell_style), 
         Paragraph(passport.passport_number or "AGB 3-5 NQ 0000125", cell_style), 
         tool_type],
        [Paragraph("Матрица / Matrix", cell_style), 
         Paragraph("Высота матрицы / Imp Depth", cell_style), 
         Paragraph("Промывочные отверстия / Waterways", cell_style), 
         Paragraph("Дата производства / Production date", cell_style)],
        [Paragraph(nomenclature.matrix or "3-5", cell_style), 
         Paragraph(nomenclature.height or "12 мм", cell_style), 
         Paragraph("8 mm", cell_style), 
         "2025"],
        [Paragraph("www.almazgeobur.ru", cell_style), "", "", ""]
    ]

    # Создаем основную таблицу (без дублирующей рамки) с правильными размерами
    table = Table(passport_data, colWidths=[40*mm, 40*mm, 50*mm, 40*mm])
    table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), normal_font),
        ('FONTSIZE', (0, 0), (-1, -1), 7),  # Уменьшенный размер шрифта
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('BOX', (0, 0), (-1, -1), 0.5, colors.black),  # Внешняя рамка
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('BACKGROUND', (0, 2), (-1, 2), colors.lightgrey),
        ('SPAN', (0, 4), (3, 4)),  # Объединяем ячейку "www.almazgeobur.ru" по всей ширине
        ('ALIGN', (0, 4), (3, 4), 'CENTER'),  # Центрируем "www.almazgeobur.ru"
        ('ENCODING', (0, 0), (-1, -1), 'utf-8'),
        ('LEFTPADDING', (0, 0), (-1, -1), 2),  # Отступы для лучшего отображения
        ('RIGHTPADDING', (0, 0), (-1, -1), 2),
        ('TOPPADDING', (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
    ]))

    # Создаем таблицу с заголовком и основной таблицей
    inner_table = Table([[header_table], [Spacer(1, 8)], [table]], colWidths=[188*mm])
    inner_table.setStyle(TableStyle([
        ('BOX', (0, 0), (-1, -1), 1, colors.black),  # Общая рамка вокруг всего паспорта
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),  # Центрируем все содержимое
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),  # Вертикальное центрирование
        ('LEFTPADDING', (0, 0), (-1, -1), 5),
        ('RIGHTPADDING', (0, 0), (-1, -1), 5),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
    ]))

    story.append(inner_table)

    return story


def generate_passport_pdf(passport):
    """Генерирует PDF паспорта с правильной кодировкой и новым макетом"""
    # Используем ту же функцию, что и для массового экспорта, но с одним паспортом
    return generate_bulk_passports_pdf([passport])


def create_passport_content_without_header(passport, normal_font, normal_style):
    """Создает содержимое паспорта без заголовка для массовой выгрузки"""
    story = []
    
    # Получаем данные из БД
    nomenclature = passport.nomenclature
    if not nomenclature:
        print(f"❌ Номенклатура не найдена для паспорта {passport.passport_number}")
        return story
    
    # Создаем основную таблицу с данными паспорта
    barcode = f"AGB{nomenclature.article or '3501040'}-{passport.passport_number or '0000125'}"
    
    # Используем наименование инструмента из номенклатуры
    tool_type_name = nomenclature.name or "Буровой инструмент / Drilling tool"
    
    # Создаем стиль для ячеек с переносом текста
    cell_style = ParagraphStyle(
        'CellText',
        parent=normal_style,
        fontSize=6,  # Уменьшенный размер шрифта
        leading=7,  # Уменьшенный межстрочный интервал
        spaceBefore=0,
        spaceAfter=0,
        alignment=1,  # CENTER
    )
    
    # Данные паспорта с реальными данными из БД (согласно инструкциям) с переносом текста
    passport_data = [
        [Paragraph("Артикул / Stock Code", cell_style), 
         Paragraph("Типоразмер / Tool size", cell_style), 
         Paragraph("Серийный номер / Serial Number", cell_style), 
         Paragraph("Буровой инструмент / Tool type", cell_style)],
        [Paragraph(nomenclature.article or "3501040", cell_style), 
         Paragraph(nomenclature.matrix or "NQ", cell_style), 
         Paragraph(passport.passport_number or "AGB 3-5 NQ 0000125", cell_style), 
         Paragraph(tool_type_name, cell_style)],
        [Paragraph("Матрица / Matrix", cell_style), 
         Paragraph("Высота матрицы / Imp Depth", cell_style), 
         Paragraph("Промывочные отверстия / Waterways", cell_style), 
         Paragraph("Дата производства / Production date", cell_style)],
        [Paragraph(nomenclature.matrix or "3-5", cell_style), 
         Paragraph(nomenclature.height or "12 мм", cell_style), 
         Paragraph("8 mm", cell_style),
         "2025"],
        [Paragraph("www.almazgeobur.ru", cell_style), "", "", ""]
    ]
    
    # Создаем основную таблицу (без дублирующей рамки) с правильными размерами
    table = Table(passport_data, colWidths=[38*mm, 38*mm, 48*mm, 38*mm])
    table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), normal_font),
        ('FONTSIZE', (0, 0), (-1, -1), 6),  # Уменьшенный размер шрифта
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('BOX', (0, 0), (-1, -1), 0.5, colors.black),  # Внешняя рамка
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('BACKGROUND', (0, 2), (-1, 2), colors.lightgrey),
        ('SPAN', (0, 4), (3, 4)),  # Объединяем ячейку "www.almazgeobur.ru" по всей ширине
        ('ALIGN', (0, 4), (3, 4), 'CENTER'),  # Центрируем "www.almazgeobur.ru"
        ('ENCODING', (0, 0), (-1, -1), 'utf-8'),
        ('LEFTPADDING', (0, 0), (-1, -1), 1),  # Уменьшенные отступы
        ('RIGHTPADDING', (0, 0), (-1, -1), 1),
        ('TOPPADDING', (0, 0), (-1, -1), 1),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
    ]))
    
    story.append(table)
    return story


def generate_bulk_passports_pdf(passports):
    """Генерирует PDF с несколькими паспортами (по 4 на страницу)"""
    print(f"📄 Начинаем генерацию PDF для {len(passports)} паспортов")
    
    # Проверяем, есть ли шаблон Word для паспортов
    passport_template_path = None
    try:
        possible_paths = [
            '/app/templates/passport_template.docx',
            os.path.join(os.path.dirname(os.path.abspath(__file__)), 'templates', 'passport_template.docx'),
        ]
        for path in possible_paths:
            if os.path.exists(path):
                passport_template_path = path
                print(f"✅ Найден шаблон паспорта: {path}")
                break
    except Exception as e:
        print(f"⚠️ Ошибка при поиске шаблона паспорта: {e}")
    
    # Если есть шаблон, используем его
    if passport_template_path:
        try:
            return generate_passports_from_template(passports, passport_template_path)
        except Exception as e:
            print(f"⚠️ Ошибка при генерации из шаблона, используем стандартный метод: {e}")
            import traceback
            traceback.print_exc()
    
    # Стандартный метод генерации (4 паспорта на страницу)
    # Создаем PDF в памяти
    buffer = io.BytesIO()
    
    # Устанавливаем минимальные отступы страницы для размещения 6 паспортов без уменьшения размера
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=5*mm,
        rightMargin=5*mm,
        topMargin=3*mm,
        bottomMargin=3*mm
    )
    
    # Настраиваем шрифты
    normal_font = setup_cyrillic_fonts()
    
    # Создаем стили
    title_style, subtitle_style, normal_style = create_passport_styles(normal_font)
    
    story = []
    
    # Контактная информация
    contact_info = """ООО "Алмазгеобур" 125362, г. Москва, улица Водников, дом 2, стр. 14, оф. 11, тел.:+7 495 229 82 94
LLP "Almazgeobur" 125362, Moscow, Vodnikov Street, 2, building. 14, of. 11, tel.:+7 495 229 82 94,
e-mail: contact@almazgeobur.ru"""
    
    # Группируем паспорта по 6 на страницу
    for i in range(0, len(passports), 6):
        passport_group = passports[i:i+6]

        for j, passport in enumerate(passport_group):
            print(f"📄 Обрабатываем паспорт {j+1} в группе: {passport.passport_number}")
            
            # Создаем заголовочную таблицу с логотипом для каждого паспорта
            header_data = [[None, contact_info]]
            logo_img = create_logo_image()
            logo_added = False
            if logo_img and os.path.exists(logo_img):
                try:
                    print(f"✅ Логотип найден для массового паспорта {passport.passport_number}: {logo_img}")
                    print(f"   Размеры логотипа: 18мм × 5.4мм (как в наклейках)")
                    # Проверяем, что файл действительно является изображением
                    file_size = os.path.getsize(logo_img)
                    print(f"   Размер файла: {file_size} байт")
                    # Используем точно те же размеры, что и в наклейках
                    logo_cell = Image(logo_img, width=18*mm, height=5.4*mm)
                    print(f"   Логотип добавлен в массовый паспорт с размерами: {18*mm} × {5.4*mm} точек")
                    header_data[0][0] = logo_cell
                    logo_added = True
                    print(f"   ✅ Логотип успешно добавлен в header_data")
                except Exception as e:
                    print(f"⚠️ Ошибка при добавлении логотипа в массовый паспорт: {e}")
                    import traceback
                    traceback.print_exc()
                    # Пытаемся использовать ImageReader для более надежной загрузки
                    try:
                        from reportlab.lib.utils import ImageReader
                        logo_cell = Image(ImageReader(logo_img), width=18*mm, height=5.4*mm)
                        header_data[0][0] = logo_cell
                        logo_added = True
                        print(f"   ✅ Логотип добавлен через ImageReader")
                    except Exception as e2:
                        print(f"⚠️ Ошибка при добавлении логотипа через ImageReader: {e2}")
            
            if not logo_added:
                print(f"⚠️ Логотип не найден для массового паспорта {passport.passport_number}: {logo_img}")
                # Если логотип не найден, оставляем пустую ячейку (не добавляем текстовый fallback)
                header_data[0][0] = ""  # Пустая строка вместо None

            # Создаем таблицу ПОСЛЕ добавления логотипа (стандартные размеры для 6 паспортов)
            header_table = Table(header_data, colWidths=[30*mm, 160*mm])
            header_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), normal_font),
                ('FONTSIZE', (0, 0), (-1, -1), 5),  # Стандартный размер шрифта
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (0, 0), 'LEFT'),
                ('ALIGN', (1, 0), (1, 0), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('ENCODING', (0, 0), (-1, -1), 'utf-8'),
                ('LEFTPADDING', (0, 0), (-1, -1), 1),
                ('RIGHTPADDING', (0, 0), (-1, -1), 1),
                ('TOPPADDING', (0, 0), (-1, -1), 1),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
            ]))
            
            # Создаем содержимое паспорта без заголовка
            passport_content = create_passport_content_without_header(passport, normal_font, normal_style)
            
            # Создаем полный паспорт с заголовком и общей рамкой (стандартные размеры для 6 паспортов)
            full_passport = Table([[header_table], [Spacer(1, 2*mm)], [passport_content]], colWidths=[200*mm])
            full_passport.setStyle(TableStyle([
                ('BOX', (0, 0), (-1, -1), 0.5, colors.black),  # Более тонкая рамка
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('LEFTPADDING', (0, 0), (-1, -1), 2),
                ('RIGHTPADDING', (0, 0), (-1, -1), 2),
                ('TOPPADDING', (0, 0), (-1, -1), 2),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            ]))
            
            story.append(full_passport)

            # Добавляем минимальный интервал между паспортами для 6 паспортов
            if j < len(passport_group) - 1:
                story.append(Spacer(1, 1*mm))

        # Добавляем переход на новую страницу после каждой группы из 6 паспортов
        if i + 6 < len(passports):
            story.append(PageBreak())

    # Генерируем PDF
    print(f"🔨 Строим PDF документ...")
    doc.build(story)
    buffer.seek(0)
    
    pdf_content = buffer.getvalue()
    print(f"✅ PDF успешно сгенерирован, размер: {len(pdf_content)} байт")
    
    return pdf_content


def generate_passports_from_template(passports, template_path):
    """Генерирует DOCX паспортов из шаблона Word (5 на страницу), затем конвертирует в PDF
    
    Примечание: Для конвертации DOCX в PDF требуется docx2pdf или LibreOffice.
    Если конвертация недоступна, возвращается DOCX файл.
    """
    print(f"📄 Начинаем генерацию паспортов из шаблона Word для {len(passports)} паспортов")
    
    try:
        from docxtpl import DocxTemplate
        from docx import Document
        from docx.shared import Mm
        import tempfile
        import shutil
        import os
        
        # Создаем новый документ для размещения паспортов
        doc = Document()
        
        # Устанавливаем размеры страницы A4
        section = doc.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)
        section.top_margin = Mm(10)
        section.bottom_margin = Mm(10)
        
        # Получаем путь к логотипу
        logo_path = create_logo_image()
        
        # Контактная информация
        contact_info_ru = 'ООО "Алмазгеобур" 125362, г. Москва, улица Водников, дом 2, стр. 14, оф. 11, тел.:+7 495 229 82 94'
        contact_info_en = 'LLP "Almazgeobur" 125362, Moscow, Vodnikov Street, 2, building. 14, of. 11, tel.:+7 495 229 82 94, e-mail: contact@almazgeobur.ru'
        
        # Группируем паспорта по 6 на страницу (6 строк по 1 паспорту)
        for page_idx in range(0, len(passports), 6):
            passport_group = passports[page_idx:page_idx+6]
            
            # Создаем таблицу 6x1 для размещения 6 паспортов (6 строк, 1 столбец)
            table = doc.add_table(rows=6, cols=1)
            table.style = None
            
            # Настраиваем размеры ячеек (на всю ширину страницы минус минимальные отступы)
            passport_width = Mm(200)  # 210мм - 5мм*2 отступы
            passport_height = Mm(48)  # Примерно 297мм / 6 - минимальные отступы
            
            for row in table.rows:
                cell = row.cells[0]
                cell.width = passport_width
                
                # Устанавливаем высоту строки
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                tr = cell._element.getparent()
                trPr = tr.trPr
                if trPr is None:
                    trPr = OxmlElement('w:trPr')
                    tr.insert(0, trPr)
                trHeight = OxmlElement('w:trHeight')
                trHeight.set(qn('w:val'), str(int(passport_height * 20)))
                trHeight.set(qn('w:hRule'), 'exact')
                trPr.append(trHeight)
            
            # Заполняем таблицу паспортами из шаблона
            for idx, passport in enumerate(passport_group):
                cell = table.rows[idx].cells[0]
                
                # Рендерим шаблон для одного паспорта
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                    tmp_path = tmp_file.name
                    shutil.copy(template_path, tmp_path)
                
                template = DocxTemplate(tmp_path)
                
                # Подготавливаем данные
                nomenclature = passport.nomenclature
                production_date = "2025"
                if passport.created_at:
                    production_date = passport.created_at.strftime("%Y")
                
                context = {
                    'company_name_ru': 'ООО "Алмазгеобур"',
                    'company_name_en': 'LLP "Almazgeobur"',
                    'contact_info_ru': contact_info_ru,
                    'contact_info_en': contact_info_en,
                    'nomenclature_name': nomenclature.name if nomenclature else 'Буровой инструмент',
                    'article': nomenclature.article if nomenclature else '3501040',
                    'matrix': nomenclature.matrix if nomenclature else 'NQ',
                    'height': nomenclature.height if nomenclature else '12',
                    'waterways': '8',
                    'serial_number': passport.passport_number or 'AGB 3-5 NQ 0000125',
                    'production_date': production_date,
                    'order_number': getattr(passport, 'order_number', '') or '',
                    'website': 'www.almazgeobur.ru',
                    'date': production_date,
                }
                
                # Добавляем логотип - приоритет backend/utils/templates/logo.png
                template_logo_path = None
                # ПРИОРИТЕТ 1: backend/utils/templates/logo.png
                backend_logo_path = '/app/backend/utils/templates/logo.png'
                if os.path.exists(backend_logo_path):
                    template_logo_path = backend_logo_path
                elif logo_path and os.path.exists(logo_path):
                    template_logo_path = logo_path
                else:
                    # Fallback на старый путь
                    template_logo_path = '/app/templates/logo.png'
                
                if template_logo_path and os.path.exists(template_logo_path):
                    try:
                        from docxtpl import InlineImage
                        context['logo'] = InlineImage(template, template_logo_path, width=Mm(18), height=Mm(5.4))
                        print(f"   ✅ Логотип добавлен из шаблона: {template_logo_path}")
                    except Exception as e:
                        print(f"⚠️ Ошибка при добавлении логотипа: {e}")
                        context['logo'] = None
                else:
                    context['logo'] = None
                    print(f"⚠️ Логотип не найден: {template_logo_path}")
                
                # Рендерим шаблон
                template.render(context)
                
                # Сохраняем рендеренный шаблон
                rendered_path = tmp_path.replace('.docx', '_rendered.docx')
                template.save(rendered_path)
                
                # Копируем содержимое из рендеренного шаблона в ячейку
                rendered_doc = Document(rendered_path)
                rendered_part = rendered_doc.part
                target_part = doc.part
                
                # Определяем источник содержимого (аналогично наклейкам)
                source_cell = None
                if rendered_doc.tables and len(rendered_doc.tables) > 0:
                    source_table = rendered_doc.tables[0]
                    if len(source_table.rows) == 1 and len(source_table.columns) == 1:
                        source_cell = source_table.rows[0].cells[0]
                    elif len(source_table.rows) > 0 and len(source_table.rows[0].cells) > 0:
                        source_cell = source_table.rows[0].cells[0]
                
                if source_cell is None and rendered_doc.paragraphs:
                    # Создаем временную ячейку для параграфов
                    temp_table = rendered_doc.add_table(rows=1, cols=1)
                    temp_cell = temp_table.rows[0].cells[0]
                    for para in rendered_doc.paragraphs[:]:
                        para._element.getparent().remove(para._element)
                        temp_cell._element.append(para._element)
                    source_cell = temp_cell
                
                if source_cell:
                    # Очищаем целевую ячейку
                    cell._element.clear_content()
                    
                    # Копируем весь XML содержимого ячейки целиком
                    import xml.etree.ElementTree as ET
                    from io import BytesIO
                    
                    source_tc = source_cell._element
                    for child in list(source_tc):
                        if child.tag.endswith('}tcPr'):
                            continue
                        
                        child_xml = ET.fromstring(ET.tostring(child, encoding='unicode'))
                        
                        # Обрабатываем изображения
                        for drawing in child_xml.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'):
                            blip = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                            if blip is not None:
                                embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                if embed_id and embed_id in rendered_part.rels:
                                    image_rel = rendered_part.rels[embed_id]
                                    image_blob = image_rel.target_part.blob
                                    image_part = target_part.add_image(BytesIO(image_blob))
                                    new_embed_id = image_part.partname.relative_ref
                                    blip.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed', new_embed_id)
                        
                        cell._element.append(child_xml)
                    
                    print(f"   ✅ Содержимое паспорта скопировано из шаблона")
                else:
                    # Fallback: копируем параграфы
                    for para in rendered_doc.paragraphs:
                        new_para = cell.add_paragraph()
                        if para.alignment:
                            new_para.alignment = para.alignment
                        for run in para.runs:
                            new_run = new_para.add_run(run.text)
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            if run.font.size:
                                new_run.font.size = run.font.size
                            if run.font.name:
                                new_run.font.name = run.font.name
                
                # Удаляем временные файлы
                try:
                    os.unlink(tmp_path)
                    os.unlink(rendered_path)
                except:
                    pass
            
            # Добавляем разрыв страницы (кроме последней)
            if page_idx + 6 < len(passports):
                doc.add_page_break()
        
        # Сохраняем DOCX во временный файл
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
            tmp_docx_path = tmp_docx.name
            doc.save(tmp_docx_path)
        
        # Пытаемся конвертировать DOCX в PDF
        try:
            from docx2pdf import convert
            import io
            pdf_buffer = io.BytesIO()
            convert(tmp_docx_path, pdf_buffer)
            pdf_content = pdf_buffer.getvalue()
            os.unlink(tmp_docx_path)
            print(f"✅ PDF успешно сгенерирован из шаблона, размер: {len(pdf_content)} байт")
            return pdf_content
        except ImportError:
            print("⚠️ docx2pdf не установлен, используем стандартный метод генерации PDF")
            os.unlink(tmp_docx_path)
            raise
        except Exception as conv_err:
            print(f"⚠️ Ошибка конвертации DOCX в PDF: {conv_err}")
            print("   Используем стандартный метод генерации")
            os.unlink(tmp_docx_path)
            raise
    
    except Exception as e:
        print(f"⚠️ Ошибка при генерации из шаблона: {e}")
        import traceback
        traceback.print_exc()
        raise


def generate_stickers_pdf_reportlab(passports):
    """Генерирует PDF с наклейками через reportLab (8 наклеек на страницу: 2 столбца × 4 строки)
    
    Args:
        passports: Список паспортов для генерации наклеек
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Image, Spacer, KeepTogether
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    import io
    
    print(f"🏷️ Начинаем генерацию PDF наклеек через reportLab для {len(passports)} паспортов")
    import sys
    sys.stdout.flush()
    
    if not passports:
        print(f"⚠️ ВНИМАНИЕ: Список паспортов пустой!")
        raise ValueError("Список паспортов пустой")
    
    # Настраиваем шрифты
    normal_font = setup_cyrillic_fonts()
    styles = getSampleStyleSheet()
    normal_style = styles['Normal']
    
    print(f"📋 Обрабатываем {len(passports)} паспортов")
    sys.stdout.flush()
    
    # Создаем стили для наклеек
    sticker_title_style = ParagraphStyle(
        'StickerTitle',
        parent=normal_style,
        fontSize=10,
        fontName=normal_font,
        leading=12,
        alignment=1,  # CENTER
        spaceAfter=2,
    )
    
    sticker_text_style = ParagraphStyle(
        'StickerText',
        parent=normal_style,
        fontSize=8,
        fontName=normal_font,
        leading=10,
        alignment=1,  # CENTER
        spaceAfter=1,
    )
    
    # Размеры наклейки: 105мм × 74.25мм (2 столбца × 4 строки на A4)
    sticker_width = 105 * mm
    sticker_height = 74.25 * mm
    
    # Получаем логотип
    logo_path = create_logo_image()
    logo_img = None
    if logo_path and os.path.exists(logo_path):
        try:
            logo_img = Image(logo_path, width=18*mm, height=5.4*mm)
        except Exception as e:
            print(f"⚠️ Ошибка загрузки логотипа: {e}")
    
    # Создаем PDF в памяти
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, 
                           leftMargin=0, rightMargin=0, 
                           topMargin=0, bottomMargin=0)
    
    story = []
    
    # Группируем паспорта по 8 на страницу
    for page_idx in range(0, len(passports), 8):
        passport_group = passports[page_idx:page_idx+8]
        
        # Создаем таблицу 4×2 для размещения наклеек
        sticker_data = []
        print(f"   📄 Страница {page_idx//8 + 1}: обрабатываем {len(passport_group)} паспортов")
        sys.stdout.flush()
        
        for row_idx in range(4):
            row = []
            for col_idx in range(2):
                idx = row_idx * 2 + col_idx
                if idx < len(passport_group):
                    passport = passport_group[idx]
                    # Загружаем номенклатуру
                    nomenclature = None
                    if hasattr(passport, 'nomenclature') and passport.nomenclature:
                        nomenclature = passport.nomenclature
                    elif hasattr(passport, 'nomenclature_id') and passport.nomenclature_id:
                        # Если номенклатура не загружена, пытаемся загрузить
                        try:
                            from backend.models import VEDNomenclature
                            from backend.database import get_db
                            db = next(get_db())
                            nomenclature = db.query(VEDNomenclature).filter(VEDNomenclature.id == passport.nomenclature_id).first()
                            print(f"   📦 Загружена номенклатура для {passport.passport_number}: {nomenclature.name if nomenclature else 'None'}")
                            sys.stdout.flush()
                        except Exception as e:
                            print(f"   ⚠️ Ошибка загрузки номенклатуры: {e}")
                            sys.stdout.flush()
                    
                    print(f"   🏷️ Обрабатываем наклейку {idx+1}: {passport.passport_number}, номенклатура: {nomenclature.name if nomenclature else 'None'}")
                    sys.stdout.flush()
                    
                    # Создаем содержимое наклейки - используем простую структуру
                    sticker_elements_list = []
                    
                    # Логотип
                    if logo_img:
                        sticker_elements_list.append(logo_img)
                    
                    # Текстовое содержимое - создаем отдельные Paragraph'ы
                    sticker_elements_list.append(Paragraph('ООО "Алмазгеобур"', sticker_title_style))
                    sticker_elements_list.append(Paragraph('LLP "Almazgeobur"', sticker_text_style))
                    
                    # Название номенклатуры
                    if nomenclature and nomenclature.name:
                        sticker_elements_list.append(Paragraph(nomenclature.name, sticker_title_style))
                    
                    # Артикул
                    if nomenclature and nomenclature.article:
                        sticker_elements_list.append(Paragraph(f"Артикул: {nomenclature.article}", sticker_text_style))
                    
                    # Типоразмер
                    if nomenclature and nomenclature.matrix:
                        sticker_elements_list.append(Paragraph(f"Типоразмер: {nomenclature.matrix}", sticker_text_style))
                    
                    # Серийный номер
                    if passport.passport_number:
                        sticker_elements_list.append(Paragraph(f"Серийный номер: {passport.passport_number}", sticker_text_style))
                    
                    # Дата производства
                    production_date = "2025"
                    sticker_elements_list.append(Paragraph(f"Дата производства: {production_date}", sticker_text_style))
                    
                    # Сайт
                    sticker_elements_list.append(Paragraph("www.almazgeobur.ru", sticker_text_style))
                    
                    # В reportLab для размещения нескольких элементов используем простую структуру
                    # Создаем вложенную таблицу, но более простую - только для вертикального размещения
                    if sticker_elements_list:
                        # Разделяем элементы на изображения и текст
                        images = [e for e in sticker_elements_list if isinstance(e, Image)]
                        paragraphs = [e for e in sticker_elements_list if isinstance(e, Paragraph)]
                        
                        # Создаем данные для вложенной таблицы
                        nested_rows = []
                        
                        # Добавляем логотип, если есть
                        if images:
                            for img in images:
                                nested_rows.append([img])
                        
                        # Добавляем текстовые элементы
                        for para in paragraphs:
                            nested_rows.append([para])
                        
                        # Создаем вложенную таблицу с минимальными стилями
                        if nested_rows:
                            sticker_cell_table = Table(nested_rows, colWidths=[sticker_width-4*mm])
                            sticker_cell_table.setStyle(TableStyle([
                                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                                ('LEFTPADDING', (0, 0), (-1, -1), 2*mm),
                                ('RIGHTPADDING', (0, 0), (-1, -1), 2*mm),
                                ('TOPPADDING', (0, 0), (-1, -1), 1*mm),
                                ('BOTTOMPADDING', (0, 0), (-1, -1), 1*mm),
                                ('NOGRID', (0, 0), (-1, -1)),
                            ]))
                            
                            row.append(sticker_cell_table)
                            print(f"      ✅ Создана таблица для наклейки: {len(nested_rows)} строк (изображений: {len(images)}, текста: {len(paragraphs)})")
                            sys.stdout.flush()
                        else:
                            row.append(Paragraph("Ошибка: нет данных", sticker_text_style))
                    else:
                        print(f"      ⚠️ ВНИМАНИЕ: Нет элементов для наклейки!")
                        row.append(Paragraph("Ошибка генерации", sticker_text_style))
                else:
                    row.append("")  # Пустая ячейка
            
            sticker_data.append(row)
        
        # Создаем таблицу с наклейками
        sticker_table = Table(sticker_data, 
                             colWidths=[sticker_width, sticker_width],
                             rowHeights=[sticker_height] * 4)
        
        sticker_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        
        story.append(sticker_table)
        
        # Добавляем разрыв страницы, если есть еще паспорта
        if page_idx + 8 < len(passports):
            story.append(PageBreak())
    
    # Проверяем, что story не пустой
    if not story:
        print(f"⚠️ ВНИМАНИЕ: Story пустой! Паспортов: {len(passports)}")
        raise ValueError("Не удалось создать содержимое для PDF наклеек")
    
    print(f"📊 Story содержит {len(story)} элементов")
    for i, elem in enumerate(story):
        print(f"   Элемент {i}: {type(elem).__name__}")
        if hasattr(elem, 'data'):
            print(f"      Данные: {len(elem.data) if hasattr(elem.data, '__len__') else 'N/A'} строк")
    sys.stdout.flush()
    
    if not story:
        print(f"⚠️ КРИТИЧЕСКАЯ ОШИБКА: Story пустой перед построением PDF!")
        raise ValueError("Story пустой - нет данных для генерации PDF")
    
    # Строим PDF
    try:
        print(f"🔨 Начинаем построение PDF...")
        sys.stdout.flush()
        # Важно: очищаем buffer перед построением
        buffer.seek(0)
        buffer.truncate(0)
        
        doc.build(story)
        
        buffer.seek(0)
        pdf_content = buffer.getvalue()  # Используем getvalue() вместо read() для BytesIO
        buffer.close()
        
        # Проверяем, что это валидный PDF (должен начинаться с %PDF)
        if len(pdf_content) < 4:
            print(f"⚠️ ВНИМАНИЕ: PDF файл слишком маленький: {len(pdf_content)} байт")
            raise ValueError("PDF файл слишком маленький")
        
        # Проверяем валидность PDF
        if len(pdf_content) == 0:
            print(f"⚠️ ВНИМАНИЕ: PDF файл пустой после генерации!")
            raise ValueError("PDF файл пустой после генерации")
        
        if pdf_content[:4] != b'%PDF':
            print(f"⚠️ ВНИМАНИЕ: PDF файл невалидный! Первые 50 байт: {pdf_content[:50]}")
            raise ValueError("PDF файл невалидный - не начинается с %PDF")
        
        print(f"✅ PDF наклейки сгенерированы: {len(pdf_content)} байт (валидный PDF)")
        sys.stdout.flush()
        return pdf_content
    except Exception as e:
        print(f"❌ Ошибка при построении PDF: {e}")
        import traceback
        traceback.print_exc()
        sys.stdout.flush()
        raise


def generate_stickers_docx(passports, use_template=True):
    """Генерирует DOCX с наклейками (8 на страницу, размер 52.5x148.5 мм каждая - 2 ряда по 4)
    
    Args:
        passports: Список паспортов для генерации наклеек
        use_template: Использовать ли шаблон Word (если доступен)
    """
    # Пытаемся использовать шаблон, если он доступен
    if use_template:
        try:
            from backend.utils.sticker_template_generator import generate_stickers_from_template
            return generate_stickers_from_template(passports)
        except ImportError:
            print("⚠️ Модуль шаблонов не найден, используем стандартный метод")
        except Exception as e:
            print(f"⚠️ Ошибка при использовании шаблона: {e}")
            import traceback
            traceback.print_exc()
            print("🔄 Переключаемся на стандартный метод...")
    
    # Стандартный метод генерации
    import os
    print(f"🏷️ Начинаем генерацию DOCX наклеек для {len(passports)} паспортов")
    
    try:
        from docx import Document
        from docx.shared import Mm, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("❌ python-docx не установлен")
        raise
    
    # Создаем документ
    doc = Document()
    
    # Устанавливаем размеры страницы A4 без отступов
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(0)
    section.right_margin = Mm(0)
    section.top_margin = Mm(0)
    section.bottom_margin = Mm(0)
    
    # Размеры наклейки: 105 x 74.25 mm (2 столбца по 4 строки = 8 наклеек на странице A4)
    # A4 = 210мм × 297мм, поэтому: 210/2 = 105мм ширина, 297/4 = 74.25мм высота
    sticker_width = Mm(105)
    sticker_height = Mm(74.25)
    
    # Группируем паспорта по 8 на страницу
    for page_idx in range(0, len(passports), 8):
        passport_group = passports[page_idx:page_idx+8]
        
        # Создаем таблицу 4x2 для размещения наклеек (4 строки, 2 столбца) на весь лист
        table = doc.add_table(rows=4, cols=2)
        table.style = None
        
        # Устанавливаем ширину таблицы на весь лист (210мм)
        table.autofit = False
        table.allow_autofit = False
        
        # Убираем все отступы и границы
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Убираем границы таблицы
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Убираем отступы таблицы
        tblCellMar = OxmlElement('w:tblCellMar')
        for margin_name in ['top', 'left', 'bottom', 'right']:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), '0')
            margin.set(qn('w:type'), 'dxa')
            tblCellMar.append(margin)
        tblPr.append(tblCellMar)
        
        # Устанавливаем ширину колонок и убираем отступы
        for row_idx in range(4):
            for col_idx in range(2):
                cell = table.rows[row_idx].cells[col_idx]
                cell.width = sticker_width
                
                # Убираем все отступы в ячейках (0 мм)
                tcPr = cell._element.tcPr
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    cell._element.insert(0, tcPr)
                
                # Убираем отступы (0 мм)
                tcMar = OxmlElement('w:tcMar')
                for margin_name in ['top', 'left', 'bottom', 'right']:
                    margin = OxmlElement(f'w:{margin_name}')
                    margin.set(qn('w:w'), '0')
                    margin.set(qn('w:type'), 'dxa')
                    tcMar.append(margin)
                tcPr.append(tcMar)
                
                # Устанавливаем высоту строки
                tr = cell._element.getparent()
                trPr = tr.trPr
                if trPr is None:
                    trPr = OxmlElement('w:trPr')
                    tr.insert(0, trPr)
                trHeight = OxmlElement('w:trHeight')
                trHeight.set(qn('w:val'), str(int(sticker_height * 20)))  # Конвертируем в twips (1 mm = 20 twips)
                trHeight.set(qn('w:hRule'), 'exact')
                trPr.append(trHeight)
        
        # Заполняем таблицу наклейками (строго 8 на страницу)
        for row_idx in range(4):
            for col_idx in range(2):
                idx = row_idx * 2 + col_idx
                cell = table.rows[row_idx].cells[col_idx]
                
                # Очищаем ячейку полностью
                cell.text = ''
                # Удаляем все параграфы кроме первого
                while len(cell.paragraphs) > 1:
                    p = cell.paragraphs[-1]
                    p._element.getparent().remove(p._element)
                
                # Заполняем только если есть паспорт для этой позиции
                if idx < len(passport_group):
                    passport = passport_group[idx]
                    
                    # Получаем данные
                    nomenclature = passport.nomenclature
                    if not nomenclature:
                        continue
                    
                    # Получаем дату производства
                    production_date = "2025"
                    if passport.created_at:
                        production_date = passport.created_at.strftime("%Y")
                    
                    # Очищаем первый параграф и настраиваем его
                    p = cell.paragraphs[0]
                    p.clear()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Убираем отступы в параграфе
                    pPr = p._element.get_or_add_pPr()
                    spacing = OxmlElement('w:spacing')
                    spacing.set(qn('w:before'), '0')
                    spacing.set(qn('w:after'), '0')
                    spacing.set(qn('w:line'), '160')  # Очень компактный межстрочный интервал (80% от размера шрифта)
                    spacing.set(qn('w:lineRule'), 'auto')
                    pPr.append(spacing)
                    
                    ind = OxmlElement('w:ind')
                    ind.set(qn('w:left'), '0')
                    ind.set(qn('w:right'), '0')
                    ind.set(qn('w:firstLine'), '0')
                    pPr.append(ind)
                    
                    # Добавляем логотип (вместо надписи компании)
                    logo_img = create_logo_image()
                    if logo_img and os.path.exists(logo_img):
                        try:
                            run = p.add_run()
                            # Увеличиваем размер логотипа для лучшей видимости
                            run.add_picture(logo_img, width=Mm(18), height=Mm(5.4))
                            run.add_break()  # Перенос строки после логотипа
                            print(f"✅ Логотип добавлен в наклейку {passport.passport_number}: {logo_img}")
                        except Exception as e:
                            print(f"⚠️ Ошибка при добавлении логотипа в наклейку: {e}")
                            import traceback
                            traceback.print_exc()
                            # Fallback: если логотип не загрузился, используем текст
                            run = p.add_run("ООО \"Алмазгеобур\"\nLLP \"Almazgeobur\"\n")
                            run.font.size = Pt(6)
                            run.font.bold = True
                    else:
                        print(f"⚠️ Логотип не найден для наклейки {passport.passport_number}: {logo_img}")
                        # Fallback: если логотип не найден, используем текст
                        run = p.add_run("ООО \"Алмазгеобур\"\nLLP \"Almazgeobur\"\n")
                        run.font.size = Pt(6)
                        run.font.bold = True
                    
                    # Буровой инструмент
                    run = p.add_run(f"{nomenclature.name or 'Буровой инструмент'}\n")
                    run.font.size = Pt(5)
                    run.font.bold = True
                    
                    # Артикул и Типоразмер
                    run = p.add_run(f"Артикул / Stock Code: {nomenclature.article or '3501040'}\n")
                    run.font.size = Pt(4)
                    run = p.add_run(f"Типоразмер / Tool size: {nomenclature.matrix or 'NQ'}\n")
                    run.font.size = Pt(4)
                    
                    # Высота и Промывочные отверстия
                    run = p.add_run(f"Высота матрицы / Imp Depth: {nomenclature.height or '12'} мм\n")
                    run.font.size = Pt(4)
                    run = p.add_run(f"Промывочные отверстия / Waterways: 8 mm\n")
                    run.font.size = Pt(4)
                    
                    # Серийный номер
                    run = p.add_run(f"Серийный номер / Serial Number: {passport.passport_number or 'AGB 3-5 NQ 0000125'}\n")
                    run.font.size = Pt(4)
                    run.font.bold = True
                    
                    # Дата производства
                    run = p.add_run(f"Дата производства / Production date: {production_date}\n")
                    run.font.size = Pt(4)
        
        # Добавляем разрыв страницы (кроме последней)
        if page_idx + 8 < len(passports):
            doc.add_page_break()
    
    # Сохраняем в память
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    docx_content = buffer.getvalue()
    print(f"✅ DOCX с наклейками успешно сгенерирован, размер: {len(docx_content)} байт")
    
    return docx_content
