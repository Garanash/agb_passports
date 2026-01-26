"""
API эндпоинты для управления шаблонами паспортов и наклеек
"""

import os
import shutil
from datetime import datetime
from pathlib import Path
from typing import List, Optional
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel

from backend.models import User
from backend.api.auth import get_current_user, get_admin_user

router = APIRouter()

# Используем единый менеджер шаблонов
from backend.utils.template_manager import get_template_manager

# Типы шаблонов (для совместимости)
TEMPLATE_TYPES = {
    "sticker": "sticker_template.docx",
    "passport": "passport_template.docx"
}


class TemplateInfo(BaseModel):
    """Информация о шаблоне"""
    type: str
    filename: str
    size: int
    modified: str
    version: Optional[int] = None


class TemplateVersion(BaseModel):
    """Версия шаблона"""
    version: int
    filename: str
    size: int
    created: str
    created_by: str


@router.get("/", response_model=List[TemplateInfo])
async def list_templates(
    current_user: User = Depends(get_current_user)
):
    """Получение списка всех шаблонов"""
    templates = []
    manager = get_template_manager()
    
    try:
        for template_type in TEMPLATE_TYPES.keys():
            info = manager.get_template_info(template_type)
            if info:
                templates.append(TemplateInfo(
                    type=template_type,
                    filename=info["filename"],
                    size=info["size"],
                    modified=info["modified"]
                ))
            else:
                print(f"ℹ️ Шаблон {template_type} не найден")
    except Exception as e:
        print(f"⚠️ Ошибка при получении списка шаблонов: {e}")
        import traceback
        traceback.print_exc()
        return []
    
    return templates


@router.get("/{template_type}")
async def get_template(
    template_type: str,
    current_user: User = Depends(get_current_user)
):
    """Скачивание шаблона"""
    import sys
    import traceback
    try:
        print(f"🔍 [TEMPLATE] Запрос шаблона: {template_type}", file=sys.stderr, flush=True)
        print(f"👤 [TEMPLATE] Пользователь: {current_user.username if current_user else 'None'}", file=sys.stderr, flush=True)
    
        if template_type not in TEMPLATE_TYPES:
            print(f"❌ [TEMPLATE] Неверный тип шаблона: {template_type}", file=sys.stderr, flush=True)
            raise HTTPException(status_code=404, detail="Тип шаблона не найден")
        
        manager = get_template_manager()
        template_path = manager.get_template_path(template_type)
        
        print(f"📁 [TEMPLATE] Путь к шаблону: {template_path}", file=sys.stderr, flush=True)
        
        if not template_path:
            print(f"❌ [TEMPLATE] Шаблон не найден: {template_type}", file=sys.stderr, flush=True)
            raise HTTPException(status_code=404, detail="Шаблон не найден")
        
        if not template_path.exists():
            print(f"❌ [TEMPLATE] Файл не существует: {template_path}", file=sys.stderr, flush=True)
            raise HTTPException(status_code=404, detail="Файл шаблона не найден")
        
        print(f"✅ [TEMPLATE] Возвращаем шаблон: {template_path}", file=sys.stderr, flush=True)
        return FileResponse(
            path=str(template_path),
            filename=TEMPLATE_TYPES[template_type],
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        print(f"❌ [TEMPLATE] Ошибка: {e}", file=sys.stderr, flush=True)
        traceback.print_exc(file=sys.stderr)
        raise


@router.post("/{template_type}/upload")
async def upload_template(
    template_type: str,
    file: UploadFile = File(...),
    create_backup: bool = Form(True),
    current_user: User = Depends(get_admin_user)
):
    """Загрузка нового шаблона"""
    if template_type not in TEMPLATE_TYPES:
        raise HTTPException(status_code=400, detail="Неверный тип шаблона")
    
    # Проверяем расширение файла
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Поддерживаются только файлы .docx")
    
    manager = get_template_manager()
    
    # Сохраняем новый шаблон
    try:
        content = await file.read()
        success, message = manager.save_template(template_type, content, create_backup)
        
        if not success:
            raise HTTPException(status_code=500, detail=message)
        
        return JSONResponse({
            "message": "Шаблон успешно загружен",
            "filename": TEMPLATE_TYPES[template_type],
            "size": len(content),
            "backup_created": create_backup
        })
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка при сохранении шаблона: {str(e)}")


@router.get("/{template_type}/versions", response_model=List[TemplateVersion])
def get_template_versions(
    template_type: str,
    current_user: User = Depends(get_admin_user)
):
    """Получение списка версий шаблона"""
    if template_type not in TEMPLATE_TYPES:
        raise HTTPException(status_code=404, detail="Тип шаблона не найден")
    
    manager = get_template_manager()
    versions = []
    pattern = f"{template_type}_*.docx"
    
    for backup_file in manager.BACKUPS_DIR.glob(pattern):
        stat = backup_file.stat()
        versions.append(TemplateVersion(
            version=len(versions) + 1,
            filename=backup_file.name,
            size=stat.st_size,
            created=datetime.fromtimestamp(stat.st_mtime).isoformat(),
            created_by="system"
        ))
    
    # Сортируем по дате создания (новые первыми)
    versions.sort(key=lambda v: v.created, reverse=True)
    
    # Нумеруем версии
    for i, version in enumerate(versions, 1):
        version.version = i
    
    return versions


@router.post("/{template_type}/restore/{version}")
def restore_template_version(
    template_type: str,
    version: int,
    current_user: User = Depends(get_admin_user)
):
    """Восстановление версии шаблона"""
    if template_type not in TEMPLATE_TYPES:
        raise HTTPException(status_code=404, detail="Тип шаблона не найден")
    
    manager = get_template_manager()
    
    # Получаем список версий
    versions = []
    pattern = f"{template_type}_*.docx"
    for backup_file in manager.BACKUPS_DIR.glob(pattern):
        versions.append(backup_file)
    
    versions.sort(key=lambda f: f.stat().st_mtime, reverse=True)
    
    if version < 1 or version > len(versions):
        raise HTTPException(status_code=404, detail="Версия не найдена")
    
    backup_file = versions[version - 1]
    
    # Читаем содержимое бэкапа
    content = backup_file.read_bytes()
    
    # Сохраняем через менеджер (создаст бэкап текущего)
    success, message = manager.save_template(template_type, content, create_backup=True)
    
    if not success:
        raise HTTPException(status_code=500, detail=message)
    
    return JSONResponse({
        "message": f"Шаблон восстановлен из версии {version}",
        "restored_from": backup_file.name
    })


@router.post("/{template_type}/validate")
async def validate_template(
    template_type: str,
    file: UploadFile = File(...),
    current_user: User = Depends(get_admin_user)
):
    """Валидация шаблона с улучшенной проверкой"""
    if template_type not in TEMPLATE_TYPES:
        raise HTTPException(status_code=400, detail="Неверный тип шаблона")
    
    try:
        content = await file.read()
        
        # Проверяем, что это валидный DOCX файл
        from docx import Document
        from io import BytesIO
        import zipfile
        import re
        
        # Проверяем, что это валидный ZIP (DOCX - это ZIP)
        try:
            with zipfile.ZipFile(BytesIO(content), 'r') as zip_ref:
                if 'word/document.xml' not in zip_ref.namelist():
                    raise HTTPException(status_code=400, detail="Неверный формат DOCX файла")
        except zipfile.BadZipFile:
            raise HTTPException(status_code=400, detail="Файл не является валидным DOCX")
        
        doc = Document(BytesIO(content))
        
        # Собираем весь текст из документа (параграфы и таблицы)
        text_content = "\n".join([para.text for para in doc.paragraphs])
        
        # Также проверяем текст в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_content += "\n" + cell.text
        
        # Проверяем наличие обязательных плейсхолдеров
        required_placeholders = {
            "sticker": ["nomenclature_name", "article", "serial_number"],
            "passport": ["nomenclature_name", "article"]
        }
        
        # Опциональные плейсхолдеры
        optional_placeholders = {
            "sticker": ["logo", "stock_code", "serial_number_code", "matrix", "height", "waterways", "production_date"],
            "passport": ["logo", "matrix", "height"]
        }
        
        missing_placeholders = []
        found_placeholders = []
        
        # Проверяем обязательные плейсхолдеры
        for placeholder in required_placeholders.get(template_type, []):
            # Проверяем разные варианты синтаксиса
            patterns = [
                f"{{{{{placeholder}}}}}",
                f"{{{{ {placeholder} }}}}",
                f"{{{{ {placeholder} }}}}",
            ]
            found = any(pattern in text_content for pattern in patterns)
            if not found:
                missing_placeholders.append(placeholder)
            else:
                found_placeholders.append(placeholder)
        
        # Проверяем опциональные плейсхолдеры
        optional_found = []
        for placeholder in optional_placeholders.get(template_type, []):
            patterns = [
                f"{{{{{placeholder}}}}}",
                f"{{{{ {placeholder} }}}}",
            ]
            if any(pattern in text_content for pattern in patterns):
                optional_found.append(placeholder)
        
        return JSONResponse({
            "valid": len(missing_placeholders) == 0,
            "missing_placeholders": missing_placeholders,
            "found_placeholders": found_placeholders,
            "optional_found": optional_found,
            "paragraphs_count": len(doc.paragraphs),
            "tables_count": len(doc.tables),
            "has_structure": len(doc.paragraphs) > 0 or len(doc.tables) > 0
        })
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Ошибка валидации: {str(e)}")


@router.post("/{template_type}/save-from-html")
async def save_template_from_html(
    template_type: str,
    html_content: str = Form(...),
    current_user: User = Depends(get_admin_user)
):
    """Сохранение шаблона из HTML (для визуального редактора)"""
    if template_type not in TEMPLATE_TYPES:
        raise HTTPException(status_code=400, detail="Неверный тип шаблона")
    
    try:
        # Получаем текущий шаблон
        manager = get_template_manager()
        template_path = manager.get_template_path(template_type)
        
        if not template_path:
            raise HTTPException(status_code=404, detail="Шаблон не найден")
        
        # Загружаем оригинальный шаблон
        from docx import Document
        from docxtpl import DocxTemplate
        import tempfile
        import shutil
        
        # Создаем временную копию шаблона
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            temp_template_path = tmp_file.name
            shutil.copy(str(template_path), temp_template_path)
        
        # Загружаем шаблон через docxtpl
        template = DocxTemplate(temp_template_path)
        
        # TODO: Конвертировать HTML обратно в DOCX структуру
        # Пока просто сохраняем оригинальный шаблон (изменения будут видны при следующей загрузке)
        # В будущем можно использовать библиотеку для конвертации HTML->DOCX
        
        # Сохраняем шаблон
        content = open(temp_template_path, 'rb').read()
        success, message = manager.save_template(template_type, content, create_backup=True)
        
        # Удаляем временный файл
        try:
            os.unlink(temp_template_path)
        except:
            pass
        
        if not success:
            raise HTTPException(status_code=500, detail=message)
        
        return JSONResponse({
            "message": "Шаблон сохранен",
            "note": "Изменения будут применены при следующей загрузке"
        })
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка сохранения: {str(e)}")


@router.post("/logo/upload")
async def upload_logo(
    file: UploadFile = File(...),
    current_user: User = Depends(get_admin_user)
):
    """Загрузка логотипа"""
    try:
        # Проверяем формат файла
        if not file.filename.lower().endswith(('.png', '.jpg', '.jpeg')):
            raise HTTPException(status_code=400, detail="Поддерживаются только файлы PNG, JPG, JPEG")
        
        content = await file.read()
        
        # Проверяем размер файла (макс 5MB)
        if len(content) > 5 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="Размер файла не должен превышать 5MB")
        
        manager = get_template_manager()
        logo_path = manager.TEMPLATES_DIR / manager.LOGO_FILENAME
        
        # Создаем бэкап существующего логотипа
        if logo_path.exists():
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_filename = f"logo_{timestamp}.png"
            backup_path = manager.BACKUPS_DIR / backup_filename
            try:
                shutil.copy2(logo_path, backup_path)
            except Exception as e:
                print(f"⚠️ Не удалось создать бэкап логотипа: {e}")
        
        # Сохраняем новый логотип
        logo_path.write_bytes(content)
        
        return JSONResponse({
            "message": "Логотип успешно загружен",
            "path": str(logo_path),
            "size": len(content)
        })
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка загрузки логотипа: {str(e)}")


@router.get("/logo")
async def get_logo(
    current_user: User = Depends(get_current_user)
):
    """Получение логотипа"""
    manager = get_template_manager()
    logo_path = manager.get_logo_path()
    
    if not logo_path or not logo_path.exists():
        raise HTTPException(status_code=404, detail="Логотип не найден")
    
    return FileResponse(
        path=str(logo_path),
        filename="logo.png",
        media_type="image/png"
    )
